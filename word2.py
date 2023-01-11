
# -*- coding: utf-8 -*-
import cairo
import math
import PySimpleGUI as sg
import os
import shutil
from zipfile import ZipFile 
from openpyxl import load_workbook
import pandas as pd
import numpy as np
from copy import deepcopy
from os import listdir
from os.path import isfile, join
import subprocess, sys
import requests
from bs4 import BeautifulSoup
from IPython.display import display
import unicodedata
from selenium import webdriver
from selenium.webdriver.common.by import By
import docx
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Pt
from datetime import datetime

from openpyxl import load_workbook
from zipfile import ZipFile
from selenium import webdriver
from selenium.webdriver.common.by import By

from docx import Document

from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt

def word2(potencia_conexao, pasta_sol, num_so, n_reli_troca, n_regu_troca, n_fus_faca,n_fus_reli, SO_automatica, tem_fus_por_faca, tem_fus_por_reli, tem_regu, relig_entrada, tem_troca_reg, tem_relig, relig_subs, rec_tri, rec_mono_tri, construcao_rede, caminhoSO, numero_so, nome_so, contrucao_rede, bitola, tensao_cabo, equi_contrucao_rede, num_equip_construcao_rede, recond_mono, tipo_ponto_a_mono, ponto_a_mono,tipo_ponto_b_mono,ponto_b_mono,recond,tipo_ponto_a,ponto_a,tipo_ponto_b,ponto_b,religador_se,subestacao,religadores,kv_banco,corrente_banco,tipo_ponto_banco,num_ponto_banco,regulador,fusiveis_religadores,fusiveis_faca,alimentador, tensao_alimentador):
    def TrocarPalavra(palavra):
            orig_text = paragraph.text
            alvo = "<<"+palavra+">>"
            resultado = palavra
            new_text = str.replace(orig_text, alvo, str(dicionario_word[resultado]))
            paragraph.text = new_text
    
    def TrocarPalavra2(palavra):
            orig_text = paragraph.text
            alvo = "<<"+palavra+">>"
            resultado = palavra
            # new_text = str.replace(orig_text, alvo, CorrigirPreço(float(dicionario_word[resultado])))
            # paragraph.text = new_text
    
    def TrocarFrase(frase):
            orig_text = paragraph.text
            alvo = frase
            new_text = str.replace(orig_text, alvo, dicionario_traducao_obras_necessarias['Construcao_de_rede'])
            paragraph.text = new_text
            
    def Delete_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None
    
    
            
    dicionario_GUI = {
            'NomeSO': '',
            'ContrucaoRede' : '',
            'MM' : '',
            'TensaoCabo' : '',
            'EquipContrucaoRede' : '',
            'NumEquipContrucaoRede' : '',       
            'RecondMono' : '' ,
            'TipoPontoAMono' : '',
            'PontoAMono' : '',
            'TipoPontoBMono' : '',
            'PontoBMono' : '',
            'Recond' : '' ,
            'TipoPontoA' : '',
            'PontoA' : '',
            'TipoPontoB' : '',
            'PontoB' : '',
            'ReligadorSE' : '',
            'Subestação' : '',
            'Religadores' : '',
            'kVBanco' : '',
            'CorrenteBanco' : '',
            'TipoPontoBanco' : '',
            'NumPontoBanco' : '',
            'Regulador' : '',
            'FusiveisReligadores' : '',
            'FusiveisFaca' : '',
            'Alimentador':'',
            'TansaoAlimentador':'',}
    
    dicionario_quantidades = {
            #A maioria esta com 1, mas tem de ser corrigido para os valores reais
            'Construcao_de_rede' : float(contrucao_rede)/1000,
            'Recond_Monofasico_Trifasico' : float(recond_mono),
            'Recond_Trifasico' : float(recond),
            'Religador_SE' : 1,
            'Religador' : n_reli_troca,
            'Banco_regulador' : 1,
            'Religador_Trifasico_300k' : 1,
            'Regulador' : n_regu_troca,
            'Fusivel_Religador' : n_fus_reli,
            'Fusivel_Faca' : n_fus_faca,}
    
    dicionario_obras_necessarias = {
        'Construcao_de_rede' : construcao_rede,
        'Recond_Monofasico_Trifasico' : rec_mono_tri,
        'Recond_Trifasico' : rec_tri,
        'Religador_SE' : relig_subs,
        'Religador' : tem_relig,
        'Banco_regulador' : tem_regu,
        'Religador_Trifasico_300k' : relig_entrada,
        'Regulador' : tem_troca_reg,
        'Fusivel_Religador' : tem_fus_por_reli,
        'Fusivel_Faca' : tem_fus_por_faca,}
            
    #dicionario_Di preenchido com Sim se tem serviço e com não se não há.
    dicionario_Di = {
            'var_construcao_0' : construcao_rede,
            'var_construcao_1' : rec_mono_tri,
            'var_construcao_2' : rec_tri,
            'var_construcao_3' : relig_subs,
            'var_construcao_4' : tem_relig,
            'var_construcao_5' : tem_regu,       
            'var_construcao_6' : tem_troca_reg ,
            'var_construcao_7' : tem_regu,
            'var_construcao_8' : tem_fus_por_reli,
            'var_construcao_9' : tem_fus_por_faca,
            'Pasta' : '',
            'Subestacao_desenho' : '',
            }
    
    caminho_pasta = r"C:\Users\joaof\OneDrive\Área de Trabalho\TCC\programa\Words_testes"
    caminho_word = caminho_pasta + "\\Informacao de Acesso - Modelo MINIGERACAO.docx"
    caminho_word2 = pasta_sol + "\\Informacao de Acesso - SO"+ str(SO_automatica) +" MINIGERACAO.docx" 
    
    doc1 = docx.Document(caminho_word)
    
    dicionario_GUI['NomeSO'] = nome_so
    dicionario_GUI['ContrucaoRede'] = contrucao_rede
    dicionario_GUI['MM'] = bitola
    dicionario_GUI['TensaoCabo'] = tensao_cabo
    dicionario_GUI['EquipContrucaoRede'] = equi_contrucao_rede
    dicionario_GUI['NumEquipContrucaoRede'] = num_equip_construcao_rede
    dicionario_GUI['RecondMono'] = recond_mono
    dicionario_GUI['TipoPontoAMono'] = tipo_ponto_a_mono
    dicionario_GUI['PontoAMono'] = ponto_a_mono
    dicionario_GUI['TipoPontoBMono'] = tipo_ponto_b_mono
    dicionario_GUI['PontoBMono'] = ponto_b_mono
    dicionario_GUI['Recond'] = recond
    dicionario_GUI['TipoPontoA'] = tipo_ponto_a
    dicionario_GUI['PontoA'] = ponto_a
    dicionario_GUI['TipoPontoB'] = tipo_ponto_b
    dicionario_GUI['PontoB'] = ponto_b
    dicionario_GUI['ReligadorSE'] = religador_se
    dicionario_GUI['Subestação'] = subestacao
    dicionario_GUI['Religadores'] = religadores
    dicionario_GUI['kVBanco'] = kv_banco
    dicionario_GUI['CorrenteBanco'] = corrente_banco
    dicionario_GUI['TipoPontoBanco'] = tipo_ponto_banco
    dicionario_GUI['NumPontoBanco'] = num_ponto_banco
    dicionario_GUI['Regulador'] = regulador
    dicionario_GUI['FusiveisReligadores'] = fusiveis_religadores
    dicionario_GUI['FusiveisFaca'] = fusiveis_faca
    dicionario_GUI['Alimentador'] = alimentador
    dicionario_GUI['TansaoAlimentador'] = tensao_alimentador
    
    
    
    Tipo = 'Solar'
    
    PotUsinaModulo = potencia_conexao
    # CidadeUsina = CorreçãoNomeProprio(CorreçãoNome(dicionario_pep['CidadeUsina']))
    Alimentador = dicionario_GUI['Alimentador']
    Tensao = dicionario_GUI['TansaoAlimentador']
    ContrucaoRede = dicionario_GUI['ContrucaoRede']
    MM = dicionario_GUI['MM']
    TensaoCabo = dicionario_GUI['TensaoCabo']
    EquipContrucaoRede = dicionario_GUI['EquipContrucaoRede']
    NumEquipContrucaoRede = dicionario_GUI['NumEquipContrucaoRede']
    RecondMono = dicionario_GUI['RecondMono']
    TipoPontoAMono = dicionario_GUI['TipoPontoAMono']
    PontoAMono = dicionario_GUI['PontoAMono']
    TipoPontoBMono = dicionario_GUI['TipoPontoBMono']
    PontoBMono = dicionario_GUI['PontoBMono']
    Recond = dicionario_GUI['Recond']
    TipoPontoA = dicionario_GUI['TipoPontoA']
    PontoA = dicionario_GUI['PontoA']
    TipoPontoB = dicionario_GUI['TipoPontoB']
    PontoB = dicionario_GUI['PontoB']
    ReligadorSE = dicionario_GUI['ReligadorSE']
    Subestação = dicionario_GUI['Subestação']
    Religadores = dicionario_GUI['Religadores']
    kVBanco = dicionario_GUI['kVBanco']
    CorrenteBanco = dicionario_GUI['CorrenteBanco']
    TipoPontoBanco = dicionario_GUI['TipoPontoBanco']
    NumPontoBanco = dicionario_GUI['NumPontoBanco']
    Regulador = dicionario_GUI['Regulador']
    FusiveisReligadores = dicionario_GUI['FusiveisReligadores']
    FusiveisFaca = dicionario_GUI['FusiveisFaca']
    
    dicionario_word = {
            
            'Alimentador' : Alimentador,
            'Tensao' : Tensao,
            'ContrucaoRede' : ContrucaoRede + ' m',
            'MM' : MM,
            'TensaoCabo' : TensaoCabo,
            'EquipContrucaoRede' : EquipContrucaoRede,
            'NumEquipContrucaoRede' : NumEquipContrucaoRede,
            'RecondMono' : RecondMono ,
            'TipoPontoAMono' : TipoPontoAMono,
            'PontoAMono' : PontoAMono,
            'TipoPontoBMono' : TipoPontoBMono,
            'PontoBMono' : PontoBMono,
            'Recond' : Recond ,
            'TipoPontoA' : TipoPontoA,
            'PontoA' : PontoA,
            'TipoPontoB' : TipoPontoB,
            'PontoB' : PontoB,
            'ReligadorSE' : ReligadorSE,
            'Subestação' : Subestação,
            'Religadores' : Religadores,
            'kVBanco' : kVBanco,
            'CorrenteBanco' : CorrenteBanco,
            'TipoPontoBanco' : TipoPontoBanco,
            'NumPontoBanco' : NumPontoBanco,
            'Regulador' : Regulador,
            'FusiveisReligadores' : FusiveisReligadores,
            'FusiveisFaca' : FusiveisFaca,
        }
    
    dicionario_traducao_obras_necessarias = {
            
            'Construcao_de_rede' : f'Construção de aproximadamente {dicionario_word["ContrucaoRede"]} de rede compacta protegida, cabo {dicionario_word["MM"]} mm² - SP - {dicionario_word["TensaoCabo"]}kV - XLPE, nas proximidades do {dicionario_word["EquipContrucaoRede"]} nº {dicionario_word["NumEquipContrucaoRede"]} até a cabine de medição e proteção da unidade consumidora onde estará instalada a Geração Distribuída.',
            'Recond_Monofasico_Trifasico' : f'Recondutoramento de aproximadamente {dicionario_word["RecondMono"]} Km da rede atual monofásica do Alimentador {dicionario_word["Alimentador"]} por rede compacta protegida trifásica, cabo {dicionario_word["MM"]} mm² – SP – {dicionario_word["TensaoCabo"]} kV - XLPE, a partir {dicionario_word["TipoPontoAMono"]} nº {dicionario_word["PontoAMono"]} até as proximidades {dicionario_word["TipoPontoBMono"]} nº {dicionario_word["PontoBMono"]}.',
            'Recond_Trifasico' : f'Recondutoramento de aproximadamente {dicionario_word["Recond"]} Km da rede atual do Alimentador {dicionario_word["Alimentador"]}, por rede compacta protegida, cabo {dicionario_word["MM"]} mm² – SP – {dicionario_word["TensaoCabo"]} kV - XLPE, das proximidades {dicionario_word["TipoPontoA"]} n º {dicionario_word["PontoA"]} até as proximidades {dicionario_word["TipoPontoB"]} nº {dicionario_word["PontoB"]}.',
            'Religador_SE' : f'Substituição do Religador nº {dicionario_word["ReligadorSE"]} da subestação {dicionario_word["Subestação"]} por Religador trifásico automático com sensor de presença de tensão ambos os lados e sistema de comunicação.',
            'Religador' : f'Substituição do Religador nº {dicionario_word["Religadores"]} que se encontra no Alimentador por Religador trifásico automático com sensor de presença de tensão em ambos os lados e sistema de comunicação.',
            'Banco_regulador' : f'Instalação de Banco Regulador nº {dicionario_word["kVBanco"]}/{dicionario_word["CorrenteBanco"]} com funcionalidade de fluxo reverso nas proximidades do {dicionario_word["TipoPontoBanco"]} nº {dicionario_word["NumPontoBanco"]}.',
            'Religador_Trifasico_300k' : f'Instalação de Religador trifásico automático, com sensor de presença de tensão em ambos os lados e sistema de comunicação, no ponto de conexão da unidade consumidora onde estará instalada a Geração Distribuída.',
            'Regulador' : f'Substituição do Regulador n º {dicionario_word["Regulador"]} que se encontra no Alimentador por Regulador trifásico automático com funcionalidade de fluxo reverso.',
            'Fusivel_Religador' : f'Substituição da chave fusível n° {dicionario_word["FusiveisReligadores"]} por religador trifásico automático com sensor de presença de tensão em ambos os lados e automação.',
            'Fusivel_Faca' : f'Substituição da chave fusível n° {dicionario_word["FusiveisFaca"]} por chave seccionadora faca.',
         
            }
    
    for paragraph in doc1.paragraphs:  
        print("lap")
        
        for palavra in ["<<ContrucaoRede>>",'<<RecondMono>>','<<Recond>>','<<ReligadorSE>>' ,'<<Religadores>>' ,'<<kVBanco>>','Somente acima ou igual 300kVA','<<Regulador>>','<<FusiveisReligadores>>','<<FusiveisFaca>>']:
            def CorrespondeciaDePalavra(palavra_crua):
                dicionario_traducão = {
                        "<<ContrucaoRede>>":'Construcao_de_rede',
                        '<<RecondMono>>':'Recond_Monofasico_Trifasico',
                        '<<Recond>>':'Recond_Trifasico',
                        '<<ReligadorSE>>' :'Religador_SE',
                        '<<Religadores>>' :'Religador',
                        '<<kVBanco>>':'Banco_regulador',
                        'Somente acima ou igual 300kVA':'Religador_Trifasico_300k',
                        '<<Regulador>>':'Regulador',
                        '<<FusiveisReligadores>>':'Fusivel_Religador',
                        '<<FusiveisFaca>>':'Fusivel_Faca' 
                        }
                return dicionario_traducão[palavra_crua]
                
                
            if palavra in paragraph.text:
                if dicionario_obras_necessarias[CorrespondeciaDePalavra(palavra)] == 'Sim':
    #                    Delete_paragraph(paragraph)
                    paragraph.text = ''
                    paragraph.add_run(dicionario_traducao_obras_necessarias[CorrespondeciaDePalavra(palavra)])#+"\n")
    
                else:
                    Delete_paragraph(paragraph)
    
    
    
    
    
    
    
     
    Alimentador = dicionario_GUI['Alimentador']
    Tensao = dicionario_GUI['TansaoAlimentador']
    ContrucaoRede = dicionario_GUI['ContrucaoRede']
    MM = dicionario_GUI['MM']
    TensaoCabo = dicionario_GUI['TensaoCabo']
    EquipContrucaoRede = dicionario_GUI['EquipContrucaoRede']
    NumEquipContrucaoRede = dicionario_GUI['NumEquipContrucaoRede']
    RecondMono = dicionario_GUI['RecondMono']
    TipoPontoAMono = dicionario_GUI['TipoPontoAMono']
    PontoAMono = dicionario_GUI['PontoAMono']
    TipoPontoBMono = dicionario_GUI['TipoPontoBMono']
    PontoBMono = dicionario_GUI['PontoBMono']
    Recond = dicionario_GUI['Recond']
    TipoPontoA = dicionario_GUI['TipoPontoA']
    PontoA = dicionario_GUI['PontoA']
    TipoPontoB = dicionario_GUI['TipoPontoB']
    PontoB = dicionario_GUI['PontoB']
    ReligadorSE = dicionario_GUI['ReligadorSE']
    Subestação = dicionario_GUI['Subestação']
    Religadores = dicionario_GUI['Religadores']
    kVBanco = dicionario_GUI['kVBanco']
    CorrenteBanco = dicionario_GUI['CorrenteBanco']
    TipoPontoBanco = dicionario_GUI['TipoPontoBanco']
    NumPontoBanco = dicionario_GUI['NumPontoBanco']
    Regulador = dicionario_GUI['Regulador']
    FusiveisReligadores = dicionario_GUI['FusiveisReligadores']
    FusiveisFaca = dicionario_GUI['FusiveisFaca']
    ERDFinal = '0'
    ERD = 330.89
    
    #variaveis especificas para cada operacao
    
    #Dicionarios
    dicionario_word = {
            'Alimentador' : Alimentador,
            'Tensao' : Tensao,
            'ContrucaoRede' : ContrucaoRede + ' m',
            'MM' : MM,
            'TensaoCabo' : TensaoCabo,
            'EquipContrucaoRede' : EquipContrucaoRede,
            'NumEquipContrucaoRede' : NumEquipContrucaoRede,
            'RecondMono' : RecondMono ,
            'TipoPontoAMono' : TipoPontoAMono,
            'PontoAMono' : PontoAMono,
            'TipoPontoBMono' : TipoPontoBMono,
            'PontoBMono' : PontoBMono,
            'Recond' : Recond ,
            'TipoPontoA' : TipoPontoA,
            'PontoA' : PontoA,
            'TipoPontoB' : TipoPontoB,
            'PontoB' : PontoB,
            'ReligadorSE' : ReligadorSE,
            'Subestação' : Subestação,
            'Religadores' : Religadores,
            'kVBanco' : kVBanco,
            'CorrenteBanco' : CorrenteBanco,
            'TipoPontoBanco' : TipoPontoBanco,
            'NumPontoBanco' : NumPontoBanco,
            'Regulador' : Regulador,
            'FusiveisReligadores' : FusiveisReligadores,
            'FusiveisFaca' : FusiveisFaca,
        }
    print(dicionario_Di)
    print(dicionario_Di['Pasta'])
    print(dicionario_Di['Subestacao_desenho'])
    print(Tipo)
    print(Subestação)
    print(PotUsinaModulo)
    print(dicionario_GUI['MM'])
    print(dicionario_GUI['TensaoCabo'])
    print(dicionario_GUI['Alimentador'])
    print(dicionario_GUI['TansaoAlimentador'])
    print(dicionario_GUI['ContrucaoRede'])
    print(dicionario_GUI['EquipContrucaoRede'])
    print(dicionario_GUI['NumEquipContrucaoRede'])
    
    if dicionario_Di['var_construcao_1'] == 'Sim':
        print(dicionario_GUI['RecondMono'])
        dmono = 'Sim'
        
    if dicionario_Di['var_construcao_1'] == 'Não':
        dmono = 'Não'
        
    if dicionario_Di['var_construcao_2'] == 'Sim':            
        print(dicionario_GUI['Recond'])
        dtri ='Sim'
        
    if dicionario_Di['var_construcao_2'] == 'Não':
        dtri = 'Não'
    
    if dicionario_Di['var_construcao_3'] == 'Sim':
        ReligadorSub = []
        print(dicionario_GUI['ReligadorSE'])
        ReligadorSub_nome = str('SE'+' '+str(dicionario_GUI['ReligadorSE']))        
        print(ReligadorSub_nome)
        print(ReligadorSub_nome)
        ReligadorSub.append(ReligadorSub_nome)
        print(ReligadorSub)
    if dicionario_Di['var_construcao_3'] == 'Não':
        ReligadorSub = []
        
    if dicionario_Di['var_construcao_4'] == 'Sim':     
        print(dicionario_quantidades['Religador'])
        print(dicionario_GUI['Religadores'])
        Religador = dicionario_GUI['Religadores'].split(",")
        print(Religador)
        
    if dicionario_Di['var_construcao_4'] == 'Não':
        Religador = []
        
    if dicionario_Di['var_construcao_5'] == 'Sim':
        bancoregulador = []
        bancoregulador.append('.')
        print(bancoregulador)
        
    if dicionario_Di['var_construcao_5'] == 'Não':
        bancoregulador = []
        
    if dicionario_Di['var_construcao_7'] == 'Sim':   
        print('JOAO REGULADORLARGURA')
        print(dicionario_quantidades['Regulador'])
        print('JOAO REGULADOR')
        print(dicionario_GUI['Regulador'])
        Regulador = dicionario_GUI['Regulador'].split(",")
        print(Regulador)
        
    if dicionario_Di['var_construcao_7'] == 'Não':
        Regulador = []
        
    if dicionario_Di['var_construcao_8'] == 'Sim':   
        print('JOAO FUSIVELRELIGADORLARGURA')
        print(dicionario_quantidades['Fusivel_Religador'])
        print('JOAO FUSIVELRELIGADOR')
        print(dicionario_GUI['FusiveisReligadores'])
        Fusivelreligador = dicionario_GUI['FusiveisReligadores'].split(",")
        print(Fusivelreligador)
        
    if dicionario_Di['var_construcao_8'] == 'Não':    
        Fusivelreligador = []
        
    if dicionario_Di['var_construcao_9'] == 'Sim':             
        print('JOAO FUSIVELFACALARGURA')
        print(dicionario_quantidades['Fusivel_Faca'])
        print('JOAO FUSIVELFACA')
        print(dicionario_GUI['FusiveisFaca'])
        Fusivelfaca = dicionario_GUI['FusiveisFaca'].split(",")
        print(Fusivelfaca)
        
    if dicionario_Di['var_construcao_9'] == 'Não': 
        Fusivelfaca = []
        
    Regulador = Regulador + bancoregulador
    Fusivel = Fusivelreligador + Fusivelfaca
    Religador = Religador + ReligadorSub
    elementos = Religador + Regulador + Fusivel
    print('elementos')
    print(elementos)
    numero_de_elementos = len(elementos)
    if len(elementos) == 0:
        numero_de_elementos = 1
    print('qtdelementos')
    print(numero_de_elementos)
    distanciamento_total = 360
    distanciamento = float(distanciamento_total/numero_de_elementos)
    print('distanciamento')
    print(distanciamento)    
    filepath= caminho_word2
    doc1.save(filepath)
    os.startfile(caminho_word2)