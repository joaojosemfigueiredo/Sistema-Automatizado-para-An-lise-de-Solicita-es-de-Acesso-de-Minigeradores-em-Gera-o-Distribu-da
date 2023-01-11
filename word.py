# -*- coding: utf-8 -*-
import cairo
import math
import PySimpleGUI as sg
import os
import shutil
from zipfile import ZipFile 
from openpyxl import load_workbook
import pandas as pd
import numpy as np
from copy import deepcopy
from os import listdir
from os.path import isfile, join
import subprocess, sys
import requests
from bs4 import BeautifulSoup
from IPython.display import display
import unicodedata
from selenium import webdriver
from selenium.webdriver.common.by import By
import docx
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Pt
from datetime import datetime

from openpyxl import load_workbook
from zipfile import ZipFile
from selenium import webdriver
from selenium.webdriver.common.by import By

from docx import Document

from docx.text.paragraph import Paragraph
from docx.oxml.xmlchemy import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt

from docx import Document
from docx.shared import Inches
from PIL import Image
from svglib.svglib import svg2rlg
from reportlab.graphics import renderPDF
import fitz

def word(potencia_conexao, num_so, n_reli_troca, n_regu_troca, n_fus_faca,n_fus_reli, SO_automatica, tem_fus_por_faca, tem_fus_por_reli, tem_regu, tem_troca_reg, relig_entrada, tem_relig, relig_subs, rec_tri, rec_mono_tri, construcao_rede, caminhoSO, numero_so, nome_so, contrucao_rede, bitola, tensao_cabo, equi_contrucao_rede, num_equip_construcao_rede, recond_mono, tipo_ponto_a_mono, ponto_a_mono,tipo_ponto_b_mono,ponto_b_mono,recond,tipo_ponto_a,ponto_a,tipo_ponto_b,ponto_b,religador_se,subestacao,religadores,kv_banco,corrente_banco,tipo_ponto_banco,num_ponto_banco,regulador,fusiveis_religadores,fusiveis_faca,alimentador, tensao_alimentador, UTMX, UTMY): 
    dicionario_GUI = {
            'NomeSO': '',
            'ContrucaoRede' : '',
            'MM' : '',
            'TensaoCabo' : '',
            'EquipContrucaoRede' : '',
            'NumEquipContrucaoRede' : '',       
            'RecondMono' : '' ,
            'TipoPontoAMono' : '',
            'PontoAMono' : '',
            'TipoPontoBMono' : '',
            'PontoBMono' : '',
            'Recond' : '' ,
            'TipoPontoA' : '',
            'PontoA' : '',
            'TipoPontoB' : '',
            'PontoB' : '',
            'ReligadorSE' : '',
            'Subestação' : '',
            'Religadores' : '',
            'kVBanco' : '',
            'CorrenteBanco' : '',
            'TipoPontoBanco' : '',
            'NumPontoBanco' : '',
            'Regulador' : '',
            'FusiveisReligadores' : '',
            'FusiveisFaca' : '',
            'Alimentador':'',
            'TansaoAlimentador':'',}
    
    dicionario_quantidades = {
            #A maioria esta com 1, mas tem de ser corrigido para os valores reais
            'Construcao_de_rede' : float(contrucao_rede)/1000,
            'Recond_Monofasico_Trifasico' : float(recond_mono),
            'Recond_Trifasico' : float(recond),
            'Religador_SE' : 1,
            'Religador' : n_reli_troca,
            'Banco_regulador' : 1,
            'Religador_Trifasico_300k' : 1,
            'Regulador' : n_regu_troca,
            'Fusivel_Religador' : n_fus_reli,
            'Fusivel_Faca' : n_fus_faca,}
    
    dicionario_obras_necessarias = {
        'Construcao_de_rede' : construcao_rede,
        'Recond_Monofasico_Trifasico' : rec_mono_tri,
        'Recond_Trifasico' : rec_tri,
        'Religador_SE' : relig_subs,
        'Religador' : tem_relig,
        'Banco_regulador' : tem_regu,
        'Religador_Trifasico_300k' : tem_troca_reg,
        'Regulador' : relig_entrada,
        'Fusivel_Religador' : tem_fus_por_reli,
        'Fusivel_Faca' : tem_fus_por_faca,}
            
    #dicionario_Di preenchido com Sim se tem serviço e com não se não há.
    dicionario_Di = {
            'var_construcao_0' : construcao_rede,
            'var_construcao_1' : rec_mono_tri,
            'var_construcao_2' : rec_tri,
            'var_construcao_3' : relig_subs,
            'var_construcao_4' : tem_relig,
            'var_construcao_5' : tem_regu,       
            'var_construcao_6' : tem_troca_reg ,
            'var_construcao_7' : relig_entrada,
            'var_construcao_8' : tem_fus_por_reli,
            'var_construcao_9' : tem_fus_por_faca,
            'Pasta' : '',
            'Subestacao_desenho' : '',
            }

    dicionario_GUI['NomeSO'] = nome_so
    dicionario_GUI['ContrucaoRede'] = contrucao_rede
    dicionario_GUI['MM'] = bitola
    dicionario_GUI['TensaoCabo'] = tensao_cabo
    dicionario_GUI['EquipContrucaoRede'] = equi_contrucao_rede
    dicionario_GUI['NumEquipContrucaoRede'] = num_equip_construcao_rede
    dicionario_GUI['RecondMono'] = recond_mono
    dicionario_GUI['TipoPontoAMono'] = tipo_ponto_a_mono
    dicionario_GUI['PontoAMono'] = ponto_a_mono
    dicionario_GUI['TipoPontoBMono'] = tipo_ponto_b_mono
    dicionario_GUI['PontoBMono'] = ponto_b_mono
    dicionario_GUI['Recond'] = recond
    dicionario_GUI['TipoPontoA'] = tipo_ponto_a
    dicionario_GUI['PontoA'] = ponto_a
    dicionario_GUI['TipoPontoB'] = tipo_ponto_b
    dicionario_GUI['PontoB'] = ponto_b
    dicionario_GUI['ReligadorSE'] = religador_se
    dicionario_GUI['Subestação'] = subestacao
    dicionario_GUI['Religadores'] = religadores
    dicionario_GUI['kVBanco'] = kv_banco
    dicionario_GUI['CorrenteBanco'] = corrente_banco
    dicionario_GUI['TipoPontoBanco'] = tipo_ponto_banco
    dicionario_GUI['NumPontoBanco'] = num_ponto_banco
    dicionario_GUI['Regulador'] = regulador
    dicionario_GUI['FusiveisReligadores'] = fusiveis_religadores
    dicionario_GUI['FusiveisFaca'] = fusiveis_faca
    dicionario_GUI['Alimentador'] = alimentador
    dicionario_GUI['TansaoAlimentador'] = tensao_alimentador
    
    ##adicionando sub_desenho e pasta ao dicionario_Di:
    arquivo = r'file:///C:/Users/joaof/OneDrive/Área de Trabalho/TCC/programa atualizado/Dados Tratados/Alimentador SE.xlsx'
    
    alimentadores = dicionario_GUI['Alimentador']
    alimentadores = ''.join([i for i in alimentadores if not i.isdigit()])
    alimentadores = ''.join(e for e in alimentadores if e.isalnum())
    alimentadores.replace(" ", "")
    
    df = pd.read_excel(arquivo)
    sub = df.loc[df['Sigla']==alimentadores,'Subestação']
    sub =  " ".join(sub)
    dicionario_Di['Subestacao_desenho'] = sub
    dicionario_Di['Pasta'] = caminhoSO
    
    ############################################################
    #FUNÇÃO CRIADA POR JOÃO JOSÉ, PARA CORRIGIR POSSÍVEIS ERROS DE INSERÇÃO DAS COORDENADAS NO PEP
    def anti_alfabeto(text):
        vow = ['a', 'b', 'c', 'd', 'e', 'f', 'g', 'h', 'i', 'j', 'k', 'l', 'm', 'n', 'o', 'p', 'q', 'r', 's', 't', 'u', 'v', 'w', 'x', 'y', 'z']
        chars = []
        for i in text: #No need of the two separate loops
            if i.lower() not in vow:
                chars.append(i)
        return "".join(chars)
    ############################################################
    def Criar_Pasta_SO(fonte_de_energia,nome_da_usina,potencia,escala_pot,numero_SO,destino): #Nome da usina via GUI
        
        #Mudar o diretorio para onde quiser criar a pasta da SO 
        dirPath = destino
        
        nome_diretorio = fonte_de_energia + ' ' + nome_da_usina + ' (' + potencia + ' ' + escala_pot + ')' + ' - ' + 'SO '+ numero_SO  
        
        caminhoSO = dirPath + '/' + nome_diretorio   #Path completo do novo diretorio a ser criado! 
        
        #print(caminhoSO)    
        if not os.path.isdir(caminhoSO): #Para caso ja exista um diretorio com o mesmo nome no local-alvo
            print('\nO diretorio ainda não existe, criando um novo! \n')
            os.mkdir(caminhoSO) 
            print("adicionando os arquivos")
            nomezip = Arquivos_Para_Pasta(caminhoSO,potencia,escala_pot)
            arquivos = Deszippar(caminhoSO,numero_SO,nomezip)
        else:
            print('\nO diretorio que você quer criar ja existe, tente alterar os dados \n')         
            
        return(caminhoSO, arquivos)
        
    ############################################################    
    def Arquivos_Para_Pasta(caminhoSO,potencia,escala_pot):
        potencia = float(potencia.replace(',','.'))
        repo = join('C:/Users/joaof/OneDrive/Área de Trabalho/TCC/programa atualizado', 'Repositorio')
    
        if escala_pot == 'kW' and potencia <= 75:
            print('\n\n\nmicro')
            nomezip = '/Repositorio_SO_Micro.zip'
        else:
            nomezip = '/Repositorio_SO_Mini.zip'
        try:
            for file in os.listdir(repo):     
                dir_path = caminhoSO
                
                if os.path.exists(dir_path):
                    file_path = repo + "/" + file    
                    
                    try:
                        # move files into created directory
                        shutil.copy(file_path, dir_path)
                    except Exception:
                        print("não deu -",file)
                        pass 
        except Exception:
            pass
     
        return(nomezip)
    
    def Deszippar(direct,codigo_SO,nomezip): 
        # specifying the name of the zip file
        file = direct + nomezip
        
        # open the zip file in read mode
        with ZipFile(file, 'r') as zip: 
            # extract all files to another directory
            zip.extractall(direct)
            
        # Renomeia os arquivos do diretório especificado na variável PATH_REPLACE iniciados pela palavra armazenada em REMOVE_WORD removendo-os do nome do arquivo
        # Ex: span_teste.img -> teste.img
        
        palavra_rename = "XXXXX"
        palavre_final = codigo_SO   
        PATH_REPLACE = direct
        arquivos = []
        
        for filename in os.listdir(PATH_REPLACE):
          old_name = os.path.join(direct, filename)
          
          try:
              if filename.__contains__(palavra_rename):
                print(filename)
                new_name = old_name.replace(palavra_rename,palavre_final)
                print(new_name)
                arquivos.append(new_name)
                os.rename(os.path.join(old_name), os.path.join(new_name))
          except Exception:
              pass 
          
        os.remove(file)  
        return(arquivos)
      
        
        
    def PEPtoWORD(SO,download_dir,dicionario_obras_necessarias,dicionario_quantidades,dicionario_GUI):
        def mesPalavra(mes):
            dicionario_traducao = {
                    '01': 'Janeiro',
                    '02': 'Fevereiro',
                    '03': 'Março',
                    '04': 'Abril',
                    '05': 'Maio',
                    '06': 'Junho',
                    '07': 'Julho',
                    '08': 'Agosto',
                    '09': 'Setembro',
                    '10': 'Outrubro',
                    '11': 'Novembro',
                    '12': 'Dezembro'         
                    }
            return(dicionario_traducao[mes])
        
        def insert_paragraph_after(paragraph, text=None, style=None):
            """Insert a new paragraph after the given paragraph."""
            new_p = OxmlElement("w:p")
            paragraph._p.addnext(new_p)
            new_para = Paragraph(new_p, paragraph._parent)
            if text:
                new_para.add_run(text)
            if style is not None:
                new_para.style = style
            return new_para
        
        def EditarTabela(doc1,dicionario_total,ERDFinal) :
        
        
            doc1.tables #a list of all tables in document
            custo_total_obras= 0.0
            aux = 2
            aux2 = 10
            for obra in dicionario_total:
                if dicionario_total[obra][0] == "Sim":
                    dict_construcao = dicionario_total[obra][1]
                    #print("Retrieved value: " + doc1.tables[0].cell(aux, 0).text)
                    doc1.tables[0].cell(aux, 0).text = dict_construcao['Frase']
                    doc1.tables[0].cell(aux, 1).text = str(dict_construcao['Qnt']).replace(".",",")
                    doc1.tables[0].cell(aux, 2).text = CorrigirPreço(dict_construcao['Custo Unitario']) 
                    doc1.tables[0].cell(aux, 3).text = CorrigirPreço(dict_construcao['Custo total'])
                    #Justifica as linhas da tabela:
                    doc1.tables[0].cell(aux, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    for i in range(1,4):                  
                        doc1.tables[0].cell(aux, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    aux = aux+1
                    custo_total_obras = custo_total_obras + dict_construcao['Custo total']
                else:
                    row2=doc1.tables[0].rows[aux2]
                    row2._element.getparent().remove(row2._element)
                    aux2 = aux2-1
            
            for table in doc1.tables:
                #table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for row in table.rows:           
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if "<<Total>>" in paragraph.text:
                                paragraph.text = paragraph.text.replace("<<Total>>", CorrigirPreço(custo_total_obras-ERDFinal))
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            if "<<TotalEstimado>>" in paragraph.text:
                                paragraph.text = paragraph.text.replace("<<TotalEstimado>>", CorrigirPreço(custo_total_obras))
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            if "<<ERDFinal>>" in paragraph.text:
                                if (custo_total_obras-ERDFinal) <= 0:                                
                                    paragraph.text = paragraph.text.replace("<<ERDFinal>>", CorrigirPreço(custo_total_obras))
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                else:
                                    paragraph.text = paragraph.text.replace("<<ERDFinal>>", CorrigirPreço(ERDFinal))
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:                             
                                font = run.font
                                font.size= Pt(9)
                                
              
        def DicionarioObrasTotais(dicionario_obras_necessarias,dicionario_traducao_obras_necessarias,dicionario_quantidades,dicionario_custo_unitario):
            dicionariofinal={}
            
            for obra in dicionario_obras_necessarias:
                dicionariofinal[obra] = [dicionario_obras_necessarias[obra],{
                        'Frase' : dicionario_traducao_obras_necessarias[obra],
                        'Qnt': dicionario_quantidades[obra],
                        'Custo Unitario': dicionario_custo_unitario[obra],
                        'Custo total' : dicionario_custo_unitario[obra]*dicionario_quantidades[obra]
                        }]
            del dicionariofinal['Fusivel_Faca']
                
            return dicionariofinal
         
        def Delete_paragraph(paragraph):
            p = paragraph._element
            p.getparent().remove(p)
            p._p = p._element = None
        
        def CorrigirPreço(preço_orig):
            if preço_orig <= 0:
                preço_orig = 0
            preço = str("{:.2f}".format(preço_orig))
            preço = preço.replace(".", ",")
            if preço_orig >= 1000:
                preço = preço[:-6] + '.' + preço[-6:]
                if preço_orig >= 1000000:
                    preço = preço[:-10] + '.' + preço[-10:]
                    
            preço = "R$ " + preço
            return preço              
            
            dicionario_comparacao = {
                    'representante' : '',
                    'titular' : '',
                    'rua' : '',
                    'numero_rua' : '',
                    'cep' : '',
                    'município_representante' : '',
                    'estado_representante':'',
                    'fonte' : '',
                    'potencia' : '',
                    'modulopotencia' : '',    
                    'município' : '',
                }
            
            dict_correpondencia = {
                    'representante' : '',
                    'titular' : '',
                    'rua' : 'R. Delfino Conti',
                    'numero_rua' : 's/n',
                    'cep' : '88040900',
                    'município_representante' : 'Florianópolis',
                    'estado_representante': 'SC',
                    'fonte' : '',
                    'potencia' : '',
                    'modulopotencia' : 'kW',    
                    'município' : '',
                    }
            
            dict_pj = {
                    'representante' : '',
                    'titular' : '',
                    'rua' : 'R. Delfino Conti',
                    'numero_rua' : 's/n',
                    'cep' : '88040900',
                    'município_representante' : 'Florianópolis',
                    'estado_representante': 'SC',
                    'fonte' : '',
                    'potencia' : '',
                    'modulopotencia' : 'kW',    
                    'município' : '',
                    }
    

            
            dict_obra = {
                    'representante' : 'João José Medeiros de Figueiredo',
                    'titular' : '',
                    'rua' : 'R. Delfino Conti',
                    'numero_rua' : 's/n',
                    'cep' : '88040900',
                    'município_representante' : 'Florianópolis',
                    'estado_representante': 'SC',
                    'fonte' : 'UFV',
                    'potencia' : '1000',
                    'modulopotencia' : 'kW',    
                    'município' : 'Florianópolis',       
            
                    }
           
            dict_titular = {
                    'representante' : 'João José Medeiros de Figueiredo',
                    'titular' : 'João José Medeiros de Figueiredo',
                    'rua' : 'R. Delfino Conti',
                    'numero_rua' : 's/n',
                    'cep' : '88040900',
                    'município_representante' : 'Florianópolis',
                    'estado_representante': 'SC',
                    'fonte' : 'UFV',
                    'potencia' : '1000',
                    'modulopotencia' : 'kW',    
                    'município' : 'Florianópolis',  
                    }
            
            dict_total = {'correspondencia' : dict_correpondencia,
                          'pj': dict_pj,
                          'obra' : dict_obra,
                          'titular' : dict_titular              
                          }
            
            ditnovo = {
                    'representante' : '',
                    'titular' : '',
                    'rua' : '',
                    'numero_rua' : '',
                    'cep' : '',
                    'município_representante' : '',
                    'estado_representante':'',
                    'fonte' : '',
                    'potencia' : '',
                    'modulopotencia' : '',    
                    'município' : '',       
                    }
            
            
            #Varredura dos dicionarios, em busca dos correspondentes
            for keys2 in dicionario_comparacao: #represntenate,rua...  
                for keys1 in dict_total:# dict_correpondencia,dict_obra,dict_titular  
                    auxnovo = dict_total[keys1][keys2] #dict_total[dict_obra]['representante']
                    print(auxnovo)
                    if ditnovo[keys2] == '' and auxnovo != None:
                        print('passou\n')
                        ditnovo[keys2] = auxnovo
            
            if ditnovo['rua'][0:4] == ('Rua ' or 'rua ' or 'RUA '):
               ditnovo['rua'] =  ditnovo['rua'].replace(ditnovo['rua'][0:4],"")
            
            #Variaveis
            data = datetime.today().strftime('%d')
            mes = mesPalavra(datetime.today().strftime('%m'))
            ano = '20'+datetime.today().strftime('%y')
            if ditnovo['representante'] != '':
                representante = ditnovo['representante']
            elif ditnovo['titular'] !='':
                representante = ditnovo['titular']
            else:    
                representante = '---Sem Nome informado---'
                
            rua = ditnovo['rua']
            numero_rua = ditnovo['numero_rua']
            cep = ditnovo['cep']
            município_representante = ditnovo['município_representante']
            fonte = ditnovo['fonte']
            numero_so = SO
            potencia = ditnovo['potencia'].replace('.','').replace(',','.')
            potencia = math.ceil(float(potencia))
            potencia = str(potencia).replace('.','')
            print('POTENCIA JOAO')
            print(potencia)
            modulopotencia = ditnovo['modulopotencia']
            município = ditnovo['município']
            PotUsinaModulo = 0
            ##modulo da usina!, modulo potencia
            dit_final = {
                    'Dia' : datetime.today().strftime('%d'),
                    'Mes' : mesPalavra(datetime.today().strftime('%m')),
                    'Ano' : '20'+datetime.today().strftime('%y'),
                    'Tratamento' : '',
                    'RepresentanteLegal' : representante,
                    'Endereco' : rua,
                    'NumEndereco' : numero_rua,
                    'CEP' : cep,
                    'CidadeEndereco' : município_representante,
                    'EstadoEndereco' : ditnovo['estado_representante'],
                    'PrezadoTratamento' : '',
                    'Tipo' : fonte,
                    'NomeUsina' : '',
                    'NumSO' : numero_so,
                    'PotUsina' : potencia,
                    'CidadeUsina': município,
                    'Alimentador' : '',
                    'Tensao' : '',
                    'ContrucaoRede' : '' + 'm',
                    'MM' : '',
                    'TensaoCabo' : '',
                    'EquipContrucaoRede' : '',
                    'NumEquipContrucaoRede' : '',
                    'RecondMono' : '' + 'm',
                    'TipoPontoAMono' : '',
                    'PontoAMono' : '',
                    'TipoPontoBMono' : '',
                    'PontoBMono' : '',
                    'Recond' : '' + 'm',
                    'TipoPontoA' : '',
                    'PontoA' : '',
                    'TipoPontoB' : '',
                    'PontoB' : '',
                    'ReligadorSE' : '',
                    'Subestação' : '',
                    'Religadores' : '',
                    'kVBanco' : '',
                    'CorrenteBanco' : '',
                    'TipoPontoBanco' : '',
                    'NumPontoBanco' : '',
                    'Regulador' : '',
                    'FusiveisReligadores' : '',
                    'FusiveisFaca' : '',
                    'UTM_X' :'745059.16',
                    'UTM_Y' : '6944633.11',
                    'ERDFinal' : 330.82*potencia_conexao,       
                    }
            
            return (dit_final)
                
        def a_o(caminho_pasta,dicionario_word,download_dir,dicionario_obras_necessarias,dicionario_quantidades,dicionario_GUI): 
            
            def TrocarPalavra(palavra):
                    orig_text = paragraph.text
                    alvo = "<<"+palavra+">>"
                    resultado = palavra
                    new_text = str.replace(orig_text, alvo, str(dicionario_word[resultado]))
                    paragraph.text = new_text
            
            def TrocarPalavra2(palavra):
                    orig_text = paragraph.text
                    alvo = "<<"+palavra+">>"
                    resultado = palavra
                    new_text = str.replace(orig_text, alvo, CorrigirPreço(float(dicionario_word[resultado])))
                    paragraph.text = new_text
            
        
            
            def TrocarFrase(frase):
                    orig_text = paragraph.text
                    alvo = frase
                    new_text = str.replace(orig_text, alvo, dicionario_traducao_obras_necessarias['Construcao_de_rede'])
                    paragraph.text = new_text
                    
            
            #Dicionarios contendo as frases padrão
            dicionario_traducao_obras_necessarias = {
                    
                    'Construcao_de_rede' : f'Construção de aproximadamente {dicionario_word["ContrucaoRede"]} de rede compacta protegida, cabo {dicionario_word["MM"]} mm² - SP - {dicionario_word["TensaoCabo"]}kV - XLPE, nas proximidades do {dicionario_word["EquipContrucaoRede"]} nº {dicionario_word["NumEquipContrucaoRede"]} até a cabine de medição e proteção da unidade consumidora onde estará instalada a Geração Distribuída.',
                    'Recond_Monofasico_Trifasico' : f'Recondutoramento de aproximadamente {dicionario_word["RecondMono"]} Km da rede atual monofásica do Alimentador {dicionario_word["Alimentador"]} por rede compacta protegida trifásica, cabo {dicionario_word["MM"]} mm² – SP – {dicionario_word["TensaoCabo"]} kV - XLPE, a partir {dicionario_word["TipoPontoAMono"]} nº {dicionario_word["PontoAMono"]} até as proximidades {dicionario_word["TipoPontoBMono"]} nº {dicionario_word["PontoBMono"]}.',
                    'Recond_Trifasico' : f'Recondutoramento de aproximadamente {dicionario_word["Recond"]} Km da rede atual do Alimentador {dicionario_word["Alimentador"]}, por rede compacta protegida, cabo {dicionario_word["MM"]} mm² – SP – {dicionario_word["TensaoCabo"]} kV - XLPE, das proximidades {dicionario_word["TipoPontoA"]} n º {dicionario_word["PontoA"]} até as proximidades {dicionario_word["TipoPontoB"]} nº {dicionario_word["PontoB"]}.',
                    'Religador_SE' : f'Substituição do Religador {dicionario_word["ReligadorSE"]} da subestação {dicionario_word["Subestação"]} por Religador trifásico automático com sensor de presença de tensão ambos os lados e sistema de comunicação.',
                    'Religador' : f'Substituição do Religador nº {dicionario_word["Religadores"]} que se encontra no Alimentador por Religador trifásico automático com sensor de presença de tensão em ambos os lados e sistema de comunicação.',
                    'Banco_regulador' : f'Instalação de Banco Regulador nº {dicionario_word["kVBanco"]}/{dicionario_word["CorrenteBanco"]} com funcionalidade de fluxo reverso nas proximidades do {dicionario_word["TipoPontoBanco"]} nº {dicionario_word["NumPontoBanco"]}.',
                    'Religador_Trifasico_300k' : f'Instalação de Religador trifásico automático, com sensor de presença de tensão em ambos os lados e sistema de comunicação, no ponto de conexão da unidade consumidora onde estará instalada a Geração Distribuída.',
                    'Regulador' : f'Substituição do Regulador n º {dicionario_word["Regulador"]} que se encontra no Alimentador por Regulador trifásico automático com funcionalidade de fluxo reverso.',
                    'Fusivel_Religador' : f'Substituição da chave fusível n° {dicionario_word["FusiveisReligadores"]} por religador trifásico automático com sensor de presença de tensão em ambos os lados e automação.',
                    'Fusivel_Faca' : f'Substituição da chave fusível n° {dicionario_word["FusiveisFaca"]} por chave seccionadora faca.',
                 
                    }
               
            if dicionario_GUI['TansaoAlimentador'] == '15':
                if ((dicionario_GUI['MM'] == '50') and (dicionario_GUI['TensaoCabo'] == '15')):
                    dicionario_custo_unitario = {
                        
                            'Construcao_de_rede' : 194486.74,
                            'Recond_Monofasico_Trifasico' : 194486.74,
                            'Recond_Trifasico' : 194486.74,
                            'Religador_SE' : 194666.80,
                            'Religador' : 106928.57,
                            'Banco_regulador' : 434779.45,
                            'Religador_Trifasico_300k' : 106928.57,
                            'Regulador' : 434779.45,
                            'Fusivel_Religador' : 106928.57,
                            'Fusivel_Faca' : 0,
                                        
                            }
        
                if dicionario_GUI['MM'] == '50' and dicionario_GUI['TensaoCabo'] == '25':
                    dicionario_custo_unitario = {
                        
                            'Construcao_de_rede' : 193044.94,
                            'Recond_Monofasico_Trifasico' : 193044.94,
                            'Recond_Trifasico' : 193044.94,
                            'Religador_SE' : 194666.80,
                            'Religador' : 106928.57,
                            'Banco_regulador' : 434779.45,
                            'Religador_Trifasico_300k' : 106928.57,
                            'Regulador' : 434779.45,
                            'Fusivel_Religador' : 106928.57,
                            'Fusivel_Faca' : 0,
                                        
                            }
        
                if dicionario_GUI['MM'] == '150':
                    dicionario_custo_unitario = {
                            'Construcao_de_rede' : 284825.20,
                            'Recond_Monofasico_Trifasico' : 284825.20,
                            'Recond_Trifasico' : 284825.20,
                            'Religador_SE' : 194666.80,
                            'Religador' : 106928.57,
                            'Banco_regulador' : 434779.45,
                            'Religador_Trifasico_300k' : 106928.57,
                            'Regulador' : 434779.45,
                            'Fusivel_Religador' : 106928.57,
                            'Fusivel_Faca' : 0,
                                        
                            } 
        
                if dicionario_GUI['MM'] == '185':
                    dicionario_custo_unitario = {
                            'Construcao_de_rede' : 286989.30,
                            'Recond_Monofasico_Trifasico' : 286989.30,
                            'Recond_Trifasico' : 286989.30,
                            'Religador_SE' : 194666.80,
                            'Religador' : 106928.57,
                            'Banco_regulador' : 434779.45,
                            'Religador_Trifasico_300k' : 106928.57,
                            'Regulador' : 434779.45,
                            'Fusivel_Religador' : 106928.57,
                            'Fusivel_Faca' : 0,
                                        
                            }
        
            if dicionario_GUI['TansaoAlimentador'] == '23':
                if ((dicionario_GUI['MM'] == '50') and (dicionario_GUI['TensaoCabo'] == '15')):
                    dicionario_custo_unitario = {
                            'Construcao_de_rede' : 194486.74,
                            'Recond_Monofasico_Trifasico' : 194486.74,
                            'Recond_Trifasico' : 194486.74,
                            'Religador_SE' : 194666.80,
                            'Religador' : 106928.57,
                            'Banco_regulador' : 452078.85,
                            'Religador_Trifasico_300k' : 106928.57,
                            'Regulador' : 452078.85,
                            'Fusivel_Religador' : 106928.57,
                            'Fusivel_Faca' : 0,
                                        
                            }
        
                if dicionario_GUI['MM'] == '50' and dicionario_GUI['TensaoCabo'] == '25':
                    dicionario_custo_unitario = {
                            'Construcao_de_rede' : 193044.94,
                            'Recond_Monofasico_Trifasico' : 193044.94,
                            'Recond_Trifasico' : 193044.94,
                            'Religador_SE' : 194666.80,
                            'Religador' : 106928.57,
                            'Banco_regulador' : 452078.85,
                            'Religador_Trifasico_300k' : 106928.57,
                            'Regulador' : 452078.85,
                            'Fusivel_Religador' : 106928.57,
                            'Fusivel_Faca' : 0,
                                        
                            }
        
                if dicionario_GUI['MM'] == '150':
                    dicionario_custo_unitario = {
                            'Construcao_de_rede' : 284825.20,
                            'Recond_Monofasico_Trifasico' : 284825.20,
                            'Recond_Trifasico' : 284825.20,
                            'Religador_SE' : 194666.80,
                            'Religador' : 106928.57,
                            'Banco_regulador' : 452078.85,
                            'Religador_Trifasico_300k' : 106928.57,
                            'Regulador' : 452078.85,
                            'Fusivel_Religador' : 106928.57,
                            'Fusivel_Faca' : 0,
                                        
                            } 
        
                if dicionario_GUI['MM'] == '185':
                    dicionario_custo_unitario = {
                            'Construcao_de_rede' : 286989.30,
                            'Recond_Monofasico_Trifasico' : 286989.30,
                            'Recond_Trifasico' : 286989.30,
                            'Religador_SE' : 194666.80,
                            'Religador' : 106928.57,
                            'Banco_regulador' : 452078.85,
                            'Religador_Trifasico_300k' : 106928.57,
                            'Regulador' : 452078.85,
                            'Fusivel_Religador' : 106928.57,
                            'Fusivel_Faca' : 0,
                                        
                            }
                     
            print('salvar docx com nome certo')
            print(potencia_conexao)
            print(nome_so)
            print(SO_automatica)
            Tipodoc = dicionario_word['Tipo']
            print(Tipodoc)
            caminho_word = caminho_pasta + "\\Informacao de Acesso - Modelo MINIGERACAO.docx"
            caminho_word2 = download_dir + "\\Informacao de Acesso - Modelo MINIGERACAO.docx"
            caminho_word3 = download_dir + f"\\Informacao de Acesso - {Tipodoc} {nome_so} - SO {SO_automatica} ({potencia_conexao}MW).docx"
            doc1 = docx.Document(caminho_word)

            for paragraph in doc1.paragraphs:  
                print("lap")
                if "<<Ano>>" in paragraph.text:
                    TrocarPalavra('Dia')
                    TrocarPalavra('Mes')
                    TrocarPalavra('Ano')
                    print("feito")
         
                if "<<RepresentanteLegal>>" in paragraph.text:
                    TrocarPalavra('RepresentanteLegal')
                    TrocarPalavra('Tratamento')
                    
                if "<<Endereco>>" in paragraph.text:
                    TrocarPalavra('Endereco')
                    TrocarPalavra('NumEndereco')
                    TrocarPalavra('CEP')
                    
                if "<<CidadeEndereco>>" in paragraph.text:
                    TrocarPalavra('CidadeEndereco')
                    TrocarPalavra('EstadoEndereco')
                    
                if "<<PrezadoTratamento>>" in paragraph.text:
                    TrocarPalavra('PrezadoTratamento')
        
                if "<<Tipo>>" in paragraph.text:
                    TrocarPalavra('Tipo')
                    TrocarPalavra('NomeUsina')
                    TrocarPalavra('PotUsina')        
                    
                if "<<NumSO>>" in paragraph.text:
                    TrocarPalavra('NumSO')
                    
                if "<<CidadeUsina>>" in paragraph.text:
                    TrocarPalavra('CidadeUsina')
                    
                if "<<Alimentador>>" in paragraph.text:    
                    TrocarPalavra('Alimentador')
                    TrocarPalavra('Tensao')
                    
                if "<<MM>>" in paragraph.text:    
                    TrocarPalavra('MM')
                    TrocarPalavra('TensaoCabo')
                    
                if "<<UTM_X>>" in paragraph.text:    
                    TrocarPalavra('UTM_X')
                    TrocarPalavra('UTM_Y')
                    
                if "<<PotUsina>>" in paragraph.text:    
                    TrocarPalavra('PotUsina')
                    TrocarPalavra2('ERDFinal')
                    
                if "<<ERDFinal>>" in paragraph.text:    
                    TrocarPalavra2('ERDFinal')
                
                for palavra in ["<<ContrucaoRede>>",'<<RecondMono>>','<<Recond>>','<<ReligadorSE>>' ,'<<Religadores>>' ,'<<kVBanco>>','Somente acima ou igual 300kVA','<<Regulador>>','<<FusiveisReligadores>>','<<FusiveisFaca>>']:
                    def CorrespondeciaDePalavra(palavra_crua):
                        dicionario_traducão = {
                                "<<ContrucaoRede>>":'Construcao_de_rede',
                                '<<RecondMono>>':'Recond_Monofasico_Trifasico',
                                '<<Recond>>':'Recond_Trifasico',
                                '<<ReligadorSE>>' :'Religador_SE',
                                '<<Religadores>>' :'Religador',
                                '<<kVBanco>>':'Banco_regulador',
                                'Somente acima ou igual 300kVA':'Religador_Trifasico_300k',
                                '<<Regulador>>':'Regulador',
                                '<<FusiveisReligadores>>':'Fusivel_Religador',
                                '<<FusiveisFaca>>':'Fusivel_Faca' 
                                }
                        return dicionario_traducão[palavra_crua]
                        
                        
                    if palavra in paragraph.text:
                        if dicionario_obras_necessarias[CorrespondeciaDePalavra(palavra)] == 'Sim':
        #                    Delete_paragraph(paragraph)
                            paragraph.text = ''
                            paragraph.add_run(dicionario_traducao_obras_necessarias[CorrespondeciaDePalavra(palavra)])#+"\n")
    
                        else:
                            Delete_paragraph(paragraph)
    
    
                #Adiciona o texto em negrito  
                if "Informação de Acesso – Estudo Preliminar de Viabilidade" in paragraph.text: 
                    paragraph.text = ''
                    paragraph.add_run(f"Informação de Acesso – Estudo Preliminar de Viabilidade                                          {dicionario_word['Tipo']} {dicionario_GUI['NomeSO']} ({dicionario_word['PotUsina']}) ").bold = True
          
            aux = 9
            # for key in dicionario_obras_necessarias:
            #     if dicionario_obras_necessarias[key] == 'Sim':
            #         print(key)
            #         insert_paragraph_after(doc1.paragraphs[aux],"\t• "+ dicionario_traducao_obras_necessarias[key]+"\n")
            #         doc1.paragraphs[aux].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY            
            #         doc1.paragraphs[aux].style = doc1.styles['Normal']
            #         aux = aux+1
            
            dicionario_obras_totais = DicionarioObrasTotais(dicionario_obras_necessarias,dicionario_traducao_obras_necessarias,dicionario_quantidades,dicionario_custo_unitario)
            
            if float(dicionario_word['ERDFinal']) <= 0.0:
                dicionario_word['ERDFinal'] = '0.00'
                
            EditarTabela(doc1,dicionario_obras_totais,dicionario_word['ERDFinal']) 
            
            print(potencia_conexao)
            print(nome_so)
            print(SO_automatica)
            print(dicionario_word['Tipo'])
            testefilepath = caminho_word3
            print('TESTE JOAO')
            filepath = caminho_word2
            print(filepath)
            doc1.save(testefilepath)
            os.startfile(caminho_word2)
    
            
        def Construcao_dicionarios(dicionario_pep,dicionario_GUI):
        
            def CorreçãoNomeProprio(nome):
                nomes = nome.split(' ')
                nomes = [n.capitalize() for n in nomes]
                nome = ' '.join(nomes)
                return(nome)
            
            
            def CorreçãoNome(nome):
                while nome[-1] == ' ':
                    aux = list(nome)
                    aux[-1]=''
                    nome = ''.join(aux)           
                return nome
            
            def CorrecaoPotencia(dicionario_input):
            	if float(dicionario_input.replace(',','.'))>=1000:
            		potencia = str(float(dicionario_input.replace(',','.'))/1000).replace('.',',')
            		potencia = potencia + ' MW'
            	else:
            		potencia = str(float(dicionario_input))
            		potencia = potencia + ' kW'
            	return potencia
            
            #Variaveis              
            Dia = datetime.today().strftime('%d')
            Mes = mesPalavra(datetime.today().strftime('%m'))
            Ano = '20'+datetime.today().strftime('%y')
            Tratamento = 'Ilmo Sr.'
            RepresentanteLegal = CorreçãoNomeProprio(dicionario_pep['RepresentanteLegal'])
            Endereco = CorreçãoNomeProprio(dicionario_pep['Endereco'])
            NumEndereco = dicionario_pep['NumEndereco']
            CEP = dicionario_pep['CEP']
            CidadeEndereco = CorreçãoNomeProprio(CorreçãoNome(dicionario_pep['CidadeEndereco']))
            EstadoEndereco = dicionario_pep['EstadoEndereco']
            PrezadoTratamento = 'Prezado Senhor'
            Tipo = dicionario_pep['Tipo']
            NomeUsina = dicionario_GUI['NomeSO']
            NumSO = dicionario_pep['NumSO']
            PotUsina =  CorrecaoPotencia(dicionario_pep['PotUsina'])
            
            PotUsinaModulo = float(dicionario_pep['PotUsina'])
            CidadeUsina = CorreçãoNomeProprio(CorreçãoNome(dicionario_pep['CidadeUsina']))
            Alimentador = dicionario_GUI['Alimentador']
            Tensao = dicionario_GUI['TansaoAlimentador']
            ContrucaoRede = dicionario_GUI['ContrucaoRede']
            MM = dicionario_GUI['MM']
            TensaoCabo = dicionario_GUI['TensaoCabo']
            EquipContrucaoRede = dicionario_GUI['EquipContrucaoRede']
            NumEquipContrucaoRede = dicionario_GUI['NumEquipContrucaoRede']
            RecondMono = dicionario_GUI['RecondMono']
            TipoPontoAMono = dicionario_GUI['TipoPontoAMono']
            PontoAMono = dicionario_GUI['PontoAMono']
            TipoPontoBMono = dicionario_GUI['TipoPontoBMono']
            PontoBMono = dicionario_GUI['PontoBMono']
            Recond = dicionario_GUI['Recond']
            TipoPontoA = dicionario_GUI['TipoPontoA']
            PontoA = dicionario_GUI['PontoA']
            TipoPontoB = dicionario_GUI['TipoPontoB']
            PontoB = dicionario_GUI['PontoB']
            ReligadorSE = dicionario_GUI['ReligadorSE']
            Subestação = dicionario_GUI['Subestação']
            Religadores = dicionario_GUI['Religadores']
            kVBanco = dicionario_GUI['kVBanco']
            CorrenteBanco = dicionario_GUI['CorrenteBanco']
            TipoPontoBanco = dicionario_GUI['TipoPontoBanco']
            NumPontoBanco = dicionario_GUI['NumPontoBanco']
            Regulador = dicionario_GUI['Regulador']
            FusiveisReligadores = dicionario_GUI['FusiveisReligadores']
            FusiveisFaca = dicionario_GUI['FusiveisFaca']
            # UTM_X = str("{:.2f}".format(float(dicionario_pep['UTM_X'].replace(',','.').replace('22J','').replace('22 J','').replace('22j','').replace('22 j','').replace(' m E','').replace(' mE','').replace('J','').replace('E','').replace(' m S','').replace(' mS','').replace('m','').replace('m E','').replace('m S','').replace(' S','').strip())))
            # UTM_Y = str("{:.2f}".format(float(dicionario_pep['UTM_Y'].replace(',','.').replace('22J','').replace('22 J','').replace('22j','').replace('22 j','').replace(' m S','').replace(' mS','').replace('J','').replace('E','').replace(' m E','').replace(' mE','').replace('m','').replace('m S','').replace('m E','').replace(' S','').strip())))
            UTM_X = UTMX
            UTM_Y = UTMY
            UTM_X = str(anti_alfabeto(UTM_X))
            UTM_Y = str(anti_alfabeto(UTM_Y))
            ERDFinal = '0'
            ERD = 330.89
            
            #variaveis especificas para cada operacao
    
            #Dicionarios
            dicionario_word = {
                    
                    'Dia' : datetime.today().strftime('%d'),
                    'Mes' : mesPalavra(datetime.today().strftime('%m')),
                    'Ano' : '20'+datetime.today().strftime('%y'),
                    'Tratamento' : Tratamento,
                    'RepresentanteLegal' : RepresentanteLegal,
                    'Endereco' : Endereco,
                    'NumEndereco' : NumEndereco,
                    'CEP' : CEP,
                    'CidadeEndereco' : CidadeEndereco,
                    'EstadoEndereco' : EstadoEndereco,
                    'PrezadoTratamento' : PrezadoTratamento,
                    'Tipo' : Tipo,
                    'NomeUsina' : NomeUsina,
                    'NumSO' : NumSO,
                    'PotUsina' : PotUsina,
                    'CidadeUsina': CidadeUsina,
                    'Alimentador' : Alimentador,
                    'Tensao' : Tensao,
                    'ContrucaoRede' : ContrucaoRede + ' m',
                    'MM' : MM,
                    'TensaoCabo' : TensaoCabo,
                    'EquipContrucaoRede' : EquipContrucaoRede,
                    'NumEquipContrucaoRede' : NumEquipContrucaoRede,
                    'RecondMono' : RecondMono ,
                    'TipoPontoAMono' : TipoPontoAMono,
                    'PontoAMono' : PontoAMono,
                    'TipoPontoBMono' : TipoPontoBMono,
                    'PontoBMono' : PontoBMono,
                    'Recond' : Recond ,
                    'TipoPontoA' : TipoPontoA,
                    'PontoA' : PontoA,
                    'TipoPontoB' : TipoPontoB,
                    'PontoB' : PontoB,
                    'ReligadorSE' : ReligadorSE,
                    'Subestação' : Subestação,
                    'Religadores' : Religadores,
                    'kVBanco' : kVBanco,
                    'CorrenteBanco' : CorrenteBanco,
                    'TipoPontoBanco' : TipoPontoBanco,
                    'NumPontoBanco' : NumPontoBanco,
                    'Regulador' : Regulador,
                    'FusiveisReligadores' : FusiveisReligadores,
                    'FusiveisFaca' : FusiveisFaca,
                    'UTM_X' : UTM_X,
                    'UTM_Y' : UTM_Y,
                    'ERDFinal' : 330.82*potencia_conexao*1000,
                }
            
            if dicionario_Di['var_construcao_1'] == 'Sim':
                dmono = 'Sim'
                
            if dicionario_Di['var_construcao_1'] == 'Não':
                dmono = 'Não'
                
            if dicionario_Di['var_construcao_2'] == 'Sim':            
                dtri ='Sim'
                
            if dicionario_Di['var_construcao_2'] == 'Não':
                dtri = 'Não'
            
            if dicionario_Di['var_construcao_3'] == 'Sim':
                ReligadorSub = []
                ReligadorSub_nome = str('SE'+' '+str(dicionario_GUI['ReligadorSE']))        
                ReligadorSub.append(ReligadorSub_nome)
            if dicionario_Di['var_construcao_3'] == 'Não':
                ReligadorSub = []
                
            if dicionario_Di['var_construcao_4'] == 'Sim':     
                Religador = dicionario_GUI['Religadores'].split(",")
                
            if dicionario_Di['var_construcao_4'] == 'Não':
                Religador = []
                
            if dicionario_Di['var_construcao_5'] == 'Sim':
                bancoregulador = []
                bancoregulador.append('.')
                
            if dicionario_Di['var_construcao_5'] == 'Não':
                bancoregulador = []
                
            if dicionario_Di['var_construcao_7'] == 'Sim':   
                Regulador = dicionario_GUI['Regulador'].split(",")                
            if dicionario_Di['var_construcao_7'] == 'Não':
                Regulador = []
                
            if dicionario_Di['var_construcao_8'] == 'Sim':   
                Fusivelreligador = dicionario_GUI['FusiveisReligadores'].split(",")
                
            if dicionario_Di['var_construcao_8'] == 'Não':    
                Fusivelreligador = []
                
            if dicionario_Di['var_construcao_9'] == 'Sim':             
                Fusivelfaca = dicionario_GUI['FusiveisFaca'].split(",")                
            if dicionario_Di['var_construcao_9'] == 'Não': 
                Fusivelfaca = []
                
            Regulador = Regulador + bancoregulador
            Fusivel = Fusivelreligador + Fusivelfaca
            Religador = Religador + ReligadorSub
            elementos = Religador + Regulador + Fusivel
            print('elementos')
            print(elementos)
            numero_de_elementos = len(elementos)
            if len(elementos) == 0:
                numero_de_elementos = 1
            print('qtdelementos')
            print(numero_de_elementos)
            distanciamento_total = 360
            distanciamento = float(distanciamento_total/numero_de_elementos)
            print('distanciamento')
            print(distanciamento)    
            
            def escrita(surface,cor_a,cor_b,cor_c,tamanho_fonte,fonte,italico,negrito,posição_x,posição_y,texto):
                context = cairo.Context(surface)
                context.set_source_rgba(cor_a,cor_b,cor_c)
                context.set_font_size(tamanho_fonte)
                context.select_font_face(fonte,italico,negrito)
                context.move_to(posição_x,posição_y)
                context.text_path(texto)
                context.fill()
                context.stroke()   
                return
            
            def escritarot(surface,cor_a,cor_b,cor_c,tamanho_fonte,fonte,italico,negrito,posição_x,posição_y,texto):
                context = cairo.Context(surface)
                context.set_source_rgba(cor_a,cor_b,cor_c)
                context.set_font_size(tamanho_fonte)
                context.select_font_face(fonte,italico,negrito)
                context.move_to(posição_x,posição_y)
                context.rotate(-1.5708)
                context.text_path(texto)
                context.fill()
                context.stroke()   
                return
            
            def retangulospreenchido(surface,x_0,y_0,largura,altura,escala_x,escala_y,espessura_linha,cor_ra,cor_rb,cor_rc,cor_rd):
                context = cairo.Context(surface)
                context.rectangle(x_0, y_0, largura, altura)
                context.scale(escala_x, escala_y)
                context.set_source_rgb(1,1,1)
                context.fill_preserve()
                context.set_line_width(espessura_linha)
                context.set_source_rgba(cor_ra, cor_rb, cor_rc, cor_rd)
                context.stroke()  
                return
            
            def retangulospreenchidocinza(surface,x_0,y_0,largura,altura,escala_x,escala_y,espessura_linha,cor_ra,cor_rb,cor_rc,cor_rd):
                context = cairo.Context(surface)
                context.rectangle(x_0, y_0, largura, altura)
                context.scale(escala_x, escala_y)
                context.set_source_rgb(0.90196,0.90196,0.90196)
                context.fill_preserve()
                context.set_line_width(espessura_linha)
                context.set_source_rgba(cor_ra, cor_rb, cor_rc, cor_rd)
                context.stroke()  
                return
            
            def linhas(surface,cor_la,cor_lb,cor_lc,cor_ld,espessura,x_i,y_i,x_f,y_f):
                context = cairo.Context(surface)
                context.set_source_rgba(cor_la, cor_lb, cor_lc, cor_ld)
                context.set_line_width(espessura)
                context.set_line_cap(cairo.LINE_CAP_BUTT)
                context.move_to(x_i, y_i)
                context.line_to(x_f, y_f)
                context.stroke()
                return
            
            def linhastracejadas(surface,cor_la,cor_lb,cor_lc,cor_ld,espessura,pontoSim,pontoNão,x_i,y_i,x_f,y_f):
                context = cairo.Context(surface)
                context.set_source_rgba(cor_la, cor_lb, cor_lc, cor_ld)
                context.set_line_width(espessura)
                context.set_dash([pontoSim, pontoNão ])
                context.set_line_cap(cairo.LINE_CAP_BUTT)
                context.move_to(x_i, y_i)
                context.line_to(x_f, y_f)
                context.stroke()
                return
            
            def arcopreenchido(surface,x_0,y_0, raio,angulo_i,angulo_f,escala_x,escala_y,espessura_linha,cor_ra,cor_rb,cor_rc,cor_rd):
                context = cairo.Context(surface)
                context.arc(x_0,y_0, raio,angulo_i,angulo_f) # circulo completo - 2*(22/7)=2π
                context.scale(escala_x, escala_y)
                context.set_source_rgb(1,1,1)
                context.fill_preserve()
                context.set_line_width(espessura_linha)
                context.set_source_rgba(cor_ra, cor_rb, cor_rc, cor_rd)
                context.stroke()  
                return
            
            def arco(surface,x_0,y_0, raio,angulo_i,angulo_f,escala_x,escala_y,espessura_linha,cor_ra,cor_rb,cor_rc,cor_rd):
                context = cairo.Context(surface)
                context.arc(x_0,y_0, raio,angulo_i,angulo_f) # circulo completo - 2*(22/7)=2π
                context.scale(escala_x, escala_y)
                context.set_line_width(espessura_linha)
                context.set_source_rgba(cor_ra, cor_rb, cor_rc, cor_rd)
                context.stroke()  
                return
            
            def retangulos(surface,x_0,y_0,largura,altura,escala_x,escala_y,espessura_linha,cor_ra,cor_rb,cor_rc,cor_rd):
                context = cairo.Context(surface)
                context.rectangle(x_0, y_0, largura, altura)
                context.scale(escala_x, escala_y)
                context.set_line_width(espessura_linha)
                context.set_source_rgba(cor_ra, cor_rb, cor_rc, cor_rd)
                context.stroke()  
                return
            
            def triangulo(surface,x1,y1,x2,y2,x3,y3,cor_ra,cor_rb,cor_rc):
                context = cairo.Context(surface)
                context.move_to(x1,y1)
                context.line_to(x2,y2)
                context.line_to(x3,y3)
                context.close_path()
                context.set_source_rgb(cor_ra, cor_rb, cor_rc)
                context.fill()
                return
            
            def triangulopreenchido(surface,x1,y1,x2,y2,x3,y3,cor_ra,cor_rb,cor_rc,coral,corbl,corcl,espessura):
                context = cairo.Context(surface)
                context.move_to(x1,y1)
                context.line_to(x2,y2)
                context.line_to(x3,y3)
                context.close_path()
                context.set_source_rgb(cor_ra, cor_rb, cor_rc)
                context.fill_preserve()
                context.set_source_rgb(coral,corbl,corcl)
                context.set_line_width(espessura)
                context.stroke()
                return
            
            def losango(surface,x1,y1,x2,y2,x3,y3,x4,y4,cor_ra,cor_rb,cor_rc,coral,corbl,corcl,espessura):
                context = cairo.Context(surface)
                context.move_to(x1,y1)
                context.line_to(x2,y2)
                context.line_to(x3,y3)
                context.line_to(x4,y4)
                context.close_path()
                context.set_source_rgb(cor_ra, cor_rb, cor_rc)
                context.fill_preserve()
                context.set_source_rgb(coral,corbl,corcl)
                context.set_line_width(espessura)
                context.stroke()
                return
            
            # creating a SVG surface
            with cairo.SVGSurface(f"{dicionario_Di['Pasta']}/Resposta Consulta SO {NumSO} - Informacao de Acesso.svg", 598, 845) as surface: 
            
            ############### FIXO ##################################################################################################################
            
                retangulos(surface,3.930,238,580.791,591.314,580.791,591.314,0.001,0,0,0,1) #retangulo que define as dimensões do desenho
                retangulospreenchidocinza(surface,26.441,697.125,198.054,100.738,198.054,100.738,0.001,0,0,0,0.39215)
                linhas(surface,0,0,1,1,3.336,35.362,718.167,68.721,718.167)
                linhas(surface,0,0.5,0,1,3.336,35.362,740.842,68.721,740.842)
                linhas(surface,0.75,0,0,1,3.336,35.362,767.362,68.721,767.362)
                escrita(surface,0, 0, 1,10.34,"Arial",0,1,79.273, 722,"Instalações Celesc")
                escrita(surface,0, 0.5,0,10.34,"Arial",0,1,79.273, 745.252,"Instalações de Acessantes")
                escrita(surface,0.75, 0, 0,10.34,"Arial",0,1,79.273, 769.220,"Instalações novas a serem")
                escrita(surface,0.75, 0, 0,10.34,"Arial",0,1,79.273, 779.220,"implantadas.")
                escrita(surface,0, 0, 0,10.34,"sans-serif",0,0,30.620, 815.719,"( * ) Sistema de Proteção conforme normativa I-432.0004")
                
                escrita(surface,0, 0.69, 0.3137,9.58,"Arial",0,0,431.041,615.566,"TD")
                escrita(surface,0.75, 0, 0,8.74,"Arial",0,0,412.409,464.560,"Medidor")
                
                #ROTACIONAR 90 E ESPELHAR
                escritarot(surface,0, 0.69, 0.3137,9.58,"Arial",0,0,424.473,571.514,"Proteção")
                escritarot(surface,0, 0.69, 0.3137,9.58,"Arial",0,0,477.473,553.42,"DJ")
                escritarot(surface,0, 0.69, 0.3137,8.74,"Arial",0,0,476.435,662.058,"DJ")
                
                #arcos disjuntor 2
                arco(surface,459.303+4.864,656.599,9.81,0.666*(22/7),1.333*(22/7),459.303+4.864,656.599,0.002,0, 0.69, 0.3137, 1)
                arco(surface,463.348,664.008,2.384,0,2*(22/7),463.348,664.008,0.002,0, 0.69, 0.3137, 1)
                arco(surface,463.348,649.474,2.384,0,2*(22/7),463.348,649.474,0.002,0, 0.69, 0.3137, 1)
                #arcos disjuntor 3
                arco(surface,459.303+4.864,656.599-109.528,9.81,0.666*(22/7),1.333*(22/7),459.303+4.864,656.599-109.528,0.002,0, 0.69, 0.3137, 1)
                arco(surface,463.348,664.008-109.528,2.384,0,2*(22/7),463.348,664.008-109.528,0.002,0, 0.69, 0.3137, 1)
                arco(surface,463.348,649.474-109.528,2.384,0,2*(22/7),463.348,649.474-109.528,0.002,0, 0.69, 0.3137, 1)
                
                #Medidor
                arco(surface,423.2975,495.5695,15.4105,0,2*(22/7),423.2975,495.5695,0.002,0.75, 0, 0, 1)
                
                #TD
                arco(surface,462.403,623.316-18,12,0,2*(22/7),462.403,623.316-18,0.002,0, 0.69, 0.3137, 1)
                arco(surface,462.403,635.316-18,12,0,2*(22/7),462.403,623.316-18,0.002,0, 0.69, 0.3137, 1)
                
                #barramento do alimentador
                linhas(surface,0,0,1,1,2.290,53.789,301.781-7.403,53.789,376.59-7.403)
                
                #Linha azul e vermelha de cima depois do religador
                linhas(surface,0.75,0,0,1,1.214,462.513,329.388,462.513,422.378-7.403)
                linhas(surface,0.75,0,0,1,1.214,462.513,420.992-7,439.42+8.251,430.627-7.403)
                #seta
                linhas(surface,0,0,1,1,1.214,462.513,329.388,553.344,329.388)
                triangulo(surface,551.344,326.782,551.344,331.996,555.344,329.388,0,0,1)
            
                #detalhes fixos
                linhas(surface,0,0.69,0.3137,1,1.214,463.401,431.534,463.401,537.126)
                linhas(surface,0,0.69,0.3137,1,1.214,463.401,558.299-4.5+3.066,463.401,593.29)
                linhas(surface,0,0.69,0.3137,1,1.214,463.401,634.197-4.5,463.401,646.74)
                retangulospreenchido(surface,433.047,481.907,9.135,27.515,433.047,481.907,0.002,0.75,0,0,1)
                retangulos(surface,413.395,523.587,17.224,58.680,413.395,523.587,0.003,0,0.69,0.3137,1)
                retangulos(surface,467.983,453.539,16.331,7.338,467.983,453.539,0.002,0,0.69,0.3137,1)
                
                #seta
                linhas(surface,0,0.69,0.3137,1,1.214,463.401,457.211,478.987,457.211)
                triangulo(surface,478.987,454.004,478.987,460.418,482.987,457.211,0,0.69,0.3137)
                
                #arcos medidor
                arco(surface,463.9,480.3,4,1.5*(22/7),(22/7)/2,463.9,480.3,0.003,0.75,0,0,1)
                arco(surface,463.9,488.3,4,1.5*(22/7),(22/7)/2,463.9,488.3,0.003,0.75,0,0,1)
                linhas(surface,0,0.69,0.3137,1,1.7,453.906,566.111+0.852,463.906,566.111+0.852)
                linhas(surface,0,0.69,0.3137,1,1.7,453.906,582.111+0.852,463.906,582.111+0.852)
                arco(surface,459.51,499.49,4,(22/7)/2,1.5*(22/7),459.51,499.49,0.003,0.75,0,0,1)
                arco(surface,459.51,507.49,4,(22/7)/2,1.5*(22/7),459.51,507.49,0.003,0.75,0,0,1)
                arco(surface,449.51,499.49,4,1.5*(22/7),(22/7)/2,449.51,499.49,0.003,0.75,0,0,1)
                arco(surface,449.51,507.49,4,1.5*(22/7),(22/7)/2,449.51,507.49,0.003,0.75,0,0,1)
                
                #arcos proteção
                arco(surface,463.9,570.96,4,1.5*(22/7),(22/7)/2,463.9,570.96,0.003,0,0.69,0.3137,1)
                arco(surface,463.9,578.96,4,1.5*(22/7),(22/7)/2,463.9,578.96,0.003,0,0.69,0.3137,1)
                linhas(surface,0.75,0,0,1,1.7,453.906,475.451+0.852,463.906,475.451+0.852)
                linhas(surface,0.75,0,0,1,1.7,453.906,491.451+0.852,463.906,491.451+0.852)
                arco(surface,458.84,523.46,4,(22/7)/2,1.5*(22/7),458.84,523.46,0.003,0,0.69,0.3137,1)
                arco(surface,458.84,531.46,4,(22/7)/2,1.5*(22/7),458.84,531.46,0.003,0,0.69,0.3137,1)
                arco(surface,448.84,523.46,4,1.5*(22/7),(22/7)/2,448.84,523.46,0.003,0,0.69,0.3137,1)
                arco(surface,448.84,531.46,4,1.5*(22/7),(22/7)/2,448.84,523.46,0.003,0,0.69,0.3137,1)
                
                #linhas fixas
                linhas(surface,0,0.69,0.3137,1,1.214,483.020+0.9,457.211,490.319,457.211)
                linhas(surface,0,0.69,0.3137,1,1.214,483.92,457.211,490.319,457.211)
                linhas(surface,0,0.69,0.3137,1,1.214,490.319,452.590,490.319,461.776)
                linhas(surface,0,0.69,0.3137,1,1.214,492.73,455.287,492.73,459.378)
                linhas(surface,0,0.69,0.3137,1,1.214,429.822,527.383,448.733,527.383)
                linhas(surface,0,0.69,0.3137,1,1.214,455.358,527.383,463.203,527.383)
                linhas(surface,0,0.69,0.3137,1,1.214,429.822,574.930,463.968,574.930)
                linhas(surface,0.75,0,0,1,1.214,441.742,484.205,465.06,484.205)
                linhas(surface,0.75,0,0,1,1.214,441.742,503.383,449.587,503.383)
                linhas(surface,0.75,0,0,1,1.214,456.327,503.383,464.172,503.383)
                
                #seta preta de construção vertical
                linhas(surface,0,0,0,1,1.214,471.709,339.650,471.709,420.689)
                triangulo(surface,469.102,339.650,474.316,339.650,471.709,335.650,0,0,0)
                triangulo(surface,469.102,420.689,474.316,420.689,471.709,424.689,0,0,0)
                
            ############### VARIÁVEL  ##################################################################################################################
                
                lista = ["CGH","UHE","PCH","UTE","UTN","CGU"]
                aux = False
                for item in lista:
                    if Tipo == item:
                        # SEM INVERSOR
                        escrita(surface,0, 0.69, 0.3137,8.74,"Arial",0,0,471.085,729.619,"Sistema de Proteção (*)")
                        escrita(surface,0, 0.69, 0.3137,8.74,"Arial",0,1,479.250,761.899,"Minigeração")
                        escrita(surface,0, 0.69, 0.3137,8.74,"Arial",0,1,479.250,771.899,f"{PotUsinaModulo} kW")
                        retangulos(surface,446.842,713.685,119.204,31.092,446.842,713.685,0.001,0,0.69,0.3137,1)
                        linhas(surface,0,0.69,0.3137,1,1.214,463.401,666.051,463.401,718.388)
                        linhas(surface,0,0.69,0.3137,1,1.214,463.401,732.631+4.768,463.401,732.631+19.687+1)
                        linhas(surface,0,0.69,0.3137,1,2.290,403.366,700.961,545.202,700.961)
                        #arcominigeração
                        arco(surface,465.302-2.61,765,11.76,0,2*(22/7),285.372,765,0.002,0, 0.69, 0.3137, 1)
                        #senoides do minigerador SEM INVERSOR
                        arco(surface,458.9,765.3,3.5,(22/7),0,458.9,765.3,0.002,0,0.69,0.3137,1)
                        arco(surface,465.9,765.3,3.5,0,(22/7),465.9,765.3,0.002,0,0.69,0.3137,1)
                        #arcos disjuntor 1
                        arco(surface,459.303+4.864,656.599+70.921,9.81,0.666*(22/7),1.333*(22/7),459.303+4.864,656.599+70.921,0.002,0, 0.69, 0.3137, 1)
                        arco(surface,463.348,664.008+70.921,2.384,0,2*(22/7),463.348,664.008+70.921,0.002,0, 0.69, 0.3137, 1)
                        arco(surface,463.348,649.474+70.921,2.384,0,2*(22/7),463.348,649.474+70.921,0.002,0, 0.69, 0.3137, 1)
                        aux = True
                        
                if aux == False:
                    #COM INVERSOR
                    escrita(surface,0, 0.69, 0.3137,8.74,"Arial",0,0,293.928+176.287,550.029+163.839,"Sistema de Proteção (*)")
                    escrita(surface,0, 0.69, 0.3137,8,"sans-serif",0,0,303.357+176.287,591.795+163.839,"Inversor")
                    escrita(surface,0, 0.69, 0.3137,7.28,"sans-serif",0,0,288.477+176.287,581.492+4.683+163.839,"CA")
                    escrita(surface,0, 0.69, 0.3137,7.28,"sans-serif",0,0,273.144-1.041+176.287,598.211+5.724+163.839,"CC")
                    escrita(surface,0, 0.69, 0.3137,8.74,"Arial",0,1,301.912+176.287,628.948+163.839,"Minigeração")
                    escrita(surface,0, 0.69, 0.3137,8.74,"Arial",0,1,301.912+176.287,638.948+163.839,f"{PotUsinaModulo} kW")
                    #arcos disjuntor 4
                    arco(surface,459.303-176.314+4.864+176.287,656.599-109.528+2.141+163.839,9.81,0.666*(22/7),1.333*(22/7),459.303-176.314+4.864+176.287,656.599-109.528+2.141+163.839,0.002,0, 0.69, 0.3137, 1)
                    arco(surface,463.348-176.314+176.287,664.008-109.528+2.141+163.839,2.384,0,2*(22/7),463.348-176.314+176.287,664.008-109.528+2.141+163.839,0.002,0, 0.69, 0.3137, 1)
                    arco(surface,463.348-176.314+176.287,649.474-109.528+2.141+163.839,2.384,0,2*(22/7),463.348-176.314+176.287,649.474-109.528+2.141+163.839,0.002,0, 0.69, 0.3137, 1)
                    #arco minigeração
                    arco(surface,285.372+176.287,632.155+163.839,11.76,0,2*(22/7),285.372+176.287,632.155+163.839,0.002,0, 0.69, 0.3137, 1)
                    #senoides do minigerador COM INVERSOR
                    arco(surface,281.72+176.287,632.15+163.839,3.5,(22/7),0,281.72+176.287,632.15+163.839,0.002,0,0.69,0.3137,1)
                    arco(surface,288.72+176.287,632.15+163.839,3.5,0,(22/7),288.72+176.287,632.15+163.839,0.002,0,0.69,0.3137,1)
                    #linha verde UFV
                    linhas(surface,0,0.69,0.3137,1,1.214,285.178+1.81+176.287,501.415+0.95+163.839,285.178+1.81+176.287,538.62+0.95+163.839)
                    linhas(surface,0,0.69,0.3137,1,1.214,285.178+1.81+176.287,557.056+1.647+163.839,285.178+1.81+176.287,578.995+163.839)
                    linhas(surface,0,0.69,0.3137,1,1.214,285.178+1.81+176.287,606.679+0.95+163.839,285.178+1.81+176.287,626.485+0.95-6.591+163.839)
                    linhas(surface,0,0.69,0.3137,1,1.214,270.747+176.287,578.729+163.839,270.747+29.187+176.287,578.729+28.393+163.839)
                    linhas(surface,0,0.69,0.3137,1,2.290,226.209+176.287,518.371+163.839,368.045+176.287,518.371+163.839)
                    #retangulos
                    retangulos(surface,270.702+176.287,578.635+163.839,29.461,28.884,270.702,578.635,0.003,0,0.69,0.3137,1)
                    retangulos(surface,269.685+176.287,531.095+163.839,119.204,31.092,269.685,531.095,0.001,0,0.69,0.3137,1)
                
                escrita(surface,0.75, 0, 0,9,"Arial",0,1,475.190,369.809,"Construção de aprox. ")
                escrita(surface,0.75, 0, 0,9,"Arial",0,1,475.190,379.809,f"{dicionario_GUI['ContrucaoRede']} m REDE 3#")
                escrita(surface,0.75, 0, 0,9,"Arial",0,1,475.190,389.809,f"{dicionario_GUI['MM']} mm²-SP-{dicionario_GUI['TensaoCabo']}KV-XLPE")
                
                escrita(surface,0, 0,1,10.34,"Arial",0,1,23.553,254.227,f"SE {dicionario_Di['Subestacao_desenho']}")
                
                escrita(surface,0, 0, 1,10.34,"Arial",0,1,31.333,286.454,f"AL {dicionario_GUI['Alimentador']}")
                escrita(surface,0, 0, 1,9.58,"Arial",0,1,42.713,379.179,f"{dicionario_GUI['TansaoAlimentador']} kV")
                
                if dmono == 'Não' and dtri == 'Sim':
                    #seta preta de recondutoramento horizontal SOMENTE TRIFÁSICO OU MONOFÁSICO
                    linhas(surface,0,0,0,1,1.214,194.696,348.272,459.565,348.272)
                    triangulo(surface,194.696,345.666,194.696,350.88,190.696,348.273,0,0,0)
                    triangulo(surface,457.565,345.666,457.565,350.88,461.565,348.273,0,0,0)
                    escrita(surface,0.75, 0, 0,9,"Arial",0,1,226,360.573,f"Recondutoramento de aprox. {dicionario_GUI['Recond']} km REDE")
                    escrita(surface,0.75, 0, 0,9,"Arial",0,1,226,370.573,f"3# para REDE 3# {dicionario_GUI['MM']}mm²-SP-{dicionario_GUI['TensaoCabo']}KV-XLPE")
                    #Linha azul e vermelha de cima
                    linhas(surface,0.75,0,0,1,1.214,196.988,329.388,462.513,329.388)
                    linhas(surface,0,0,1,1,1.214,54,336.791-7.403,196.988,336.791-7.403)
                            
                elif dmono == 'Sim' and dtri == 'Não':
                    #seta preta de recondutoramento horizontal SOMENTE TRIFÁSICO OU MONOFÁSICO
                    linhas(surface,0,0,0,1,1.214,194.696,348.272,459.565,348.272)
                    triangulo(surface,194.696,345.666,194.696,350.88,190.696,348.273,0,0,0)
                    triangulo(surface,457.565,345.666,457.565,350.88,461.565,348.273,0,0,0)
                    escrita(surface,0.75, 0, 0,9,"Arial",0,1,226,360.573,f"Recondutoramento de aprox. {dicionario_GUI['RecondMono']} km REDE")
                    escrita(surface,0.75, 0, 0,9,"Arial",0,1,226,370.573,f"1# para REDE 3# {dicionario_GUI['MM']}mm²-SP-{dicionario_GUI['TensaoCabo']}KV-XLPE")
                    #Linha azul e vermelha de cima
                    linhas(surface,0.75,0,0,1,1.214,196.988,329.388,462.513,329.388)
                    linhas(surface,0,0,1,1,1.214,54,336.791-7.403,196.988,336.791-7.403)
                    
                elif dmono == 'Sim' and dtri == 'Sim':
                    #seta preta de recondutoramento horizontal OS DOIS
                    linhas(surface,0,0,0,1,1.214,194.696-121.53,348.272,459.565-121.53,348.272)
                    triangulo(surface,194.696-121.53,345.666,194.696-121.53,350.88,190.696-121.53,348.273,0,0,0)
                    triangulo(surface,457.565-121.53,345.666,457.565-121.53,350.88,461.565-121.53,348.273,0,0,0)
                    linhas(surface,0,0,0,1,1.214,459.565-115.53,348.272,459.565,348.272)
                    triangulo(surface,459.565-115.53,345.666,459.565-115.53,350.88,459.565-115.53-4,348.273,0,0,0)
                    triangulo(surface,457.565,345.666,457.565,350.88,461.565,348.273,0,0,0)
                    escrita(surface,0.75, 0, 0,9,"Arial",0,1,273.056,360.573,f"Recondutoramento de aprox. {dicionario_GUI['RecondMono']} km REDE")
                    escrita(surface,0.75, 0, 0,9,"Arial",0,1,273.056,370.573,f"1# para REDE 3# {dicionario_GUI['MM']}mm²-SP-{dicionario_GUI['TensaoCabo']}KV-XLPE")    
                    escrita(surface,0.75, 0, 0,9,"Arial",0,1,72,360.573,f"Recondutoramento de aprox. {dicionario_GUI['Recond']} km REDE")
                    escrita(surface,0.75, 0, 0,9,"Arial",0,1,72,370.573,f"3# para REDE 3# {dicionario_GUI['MM']}mm²-SP-{dicionario_GUI['TensaoCabo']}KV-XLPE")
                    #Linha azul e vermelha de cima
                    linhas(surface,0.75,0,0,1,1.214,196.988,329.388,462.513,329.388)
                    linhas(surface,0,0,1,1,1.214,54,336.791-7.403,196.988,336.791-7.403)
            
                else:
                    #Linha azul de cima
                    linhas(surface,0,0,1,1,1.214,54,336.791-7.403,462.513,336.791-7.403)
            
                aux1=False
                lista1 = ["CGH","UHE","PCH","UTE","UTN","CGU"]
                for x in lista1:
                    if Tipo == x: 
                        if dicionario_GUI['EquipContrucaoRede'] == 'Transformador':
                            #escrita
                            escrita(surface,0.75, 0, 0,6.06,"Arial",0,1,449.330,321.426,"RL")
                            escrita(surface,0, 0, 1,9.32,"sans-serif",0,0,466.387,293.213,f"TD {dicionario_GUI['NumEquipContrucaoRede']}")
                            #triangulo do trafo azul
                            linhas(surface,0,0,1,1,1.214,482.461,316.095,482.461,328.871)
                            triangulopreenchido(surface,473.854,316.095,491.068,316.095,482.461,300.095,1,1,1,0,0,1,1)
                            #Religador Vermelho de cima
                            losango(surface,31.65+421.135,547.58-218.146,41.38+421.135,537.85-218.146,51.11+421.135,547.58-218.146,41.38+421.135,557.32-218.146,1,1,1,0.75,0,0,1.8)
                            #linha tracejada religador de cima
                            linhastracejadas(surface,0.75,0,0,1,1.214,1,3,461.61,314.62,461.61,259.81)
                            escrita(surface,0.75, 0, 0,8.06,"sans-serif",0,0,360.365, 246.655,"Instalação de religador trifásico automático com")
                            escrita(surface,0.75, 0, 0,8.06,"sans-serif",0,0,360.365, 256.655,"sensor de presença de tensão em ambos os lados.")
                            aux1=True
            
                        elif dicionario_GUI['EquipContrucaoRede'] == 'Fusível' or dicionario_GUI['EquipContrucaoRede'] == 'Chave':
                            #escrita
                            escrita(surface,0.75, 0, 0,6.06,"Arial",0,1,449.330,527.562-218.151,f"FU {dicionario_GUI['NumEquipContrucaoRede']}")
                            #Fusivel vermelho
                            arcopreenchido(surface,146.5+40.723+279.5,547.71+0.874-222.981,6,0.75*(22/7),1.75*(22/7),146.5+40.723+279.5,547.71+0.874-222.981,0.004,0.75,0,0,1)
                            arcopreenchido(surface,137.97+40.723+279.5,556.31+0.874-222.981,6,1.75*(22/7),0.75*(22/7),137.97+40.723+279.5,556.31+0.874-222.981,0.004,0.75,0,0,1)
                            linhas(surface,0.75,0,0,1,1.214,133.74+40.723+279.5,560.47+0.874-222.981,150.72+40.723+279.5,543.4+0.874-222.981)
                            #linha tracejada religador de cima
                            linhastracejadas(surface,0.75,0,0,1,1.214,1,3,461.61,300,461.61,259.81)
                            escrita(surface,0.75, 0, 0,8.06,"sans-serif",0,0,360.365, 246.655,"Substituir por religador trifásico automático com")
                            escrita(surface,0.75, 0, 0,8.06,"sans-serif",0,0,360.365, 256.655,"sensor de presença de tensão em ambos os lados.")
                            aux1=True
                            
                if aux1 == False:
                    if dicionario_GUI['EquipContrucaoRede'] == 'Transformador' and PotUsinaModulo>=300:
                        #escrita
                        escrita(surface,0.75, 0, 0,6.06,"Arial",0,1,449.330,321.426,"RL")
                        escrita(surface,0, 0, 1,9.32,"sans-serif",0,0,466.387,293.213,f"TD {dicionario_GUI['NumEquipContrucaoRede']}")
                        #triangulo do trafo azul
                        linhas(surface,0,0,1,1,1.214,482.461,316.095,482.461,328.871)
                        triangulopreenchido(surface,473.854,316.095,491.068,316.095,482.461,300.095,1,1,1,0,0,1,1)
                        #Religador Vermelho de cima
                        losango(surface,31.65+421.135,547.58-218.146,41.38+421.135,537.85-218.146,51.11+421.135,547.58-218.146,41.38+421.135,557.32-218.146,1,1,1,0.75,0,0,1.8)
                        #linha tracejada religador de cima
                        linhastracejadas(surface,0.75,0,0,1,1.214,1,3,461.61,314.62,461.61,259.81)
                        escrita(surface,0.75, 0, 0,8.06,"sans-serif",0,0,360.365, 246.655,"Instalação de religador trifásico automático com")
                        escrita(surface,0.75, 0, 0,8.06,"sans-serif",0,0,360.365, 256.655,"sensor de presença de tensão em ambos os lados.")
            
                    elif dicionario_GUI['EquipContrucaoRede'] == 'Fusível' or dicionario_GUI['EquipContrucaoRede'] == 'Chave' and PotUsinaModulo>=300:
                        #escrita
                        escrita(surface,0.75, 0, 0,6.06,"Arial",0,1,449.330,527.562-218.151,f"FU {dicionario_GUI['NumEquipContrucaoRede']}")
                        #Fusivel vermelho
                        arcopreenchido(surface,146.5+40.723+279.5,547.71+0.874-222.981,6,0.75*(22/7),1.75*(22/7),146.5+40.723+279.5,547.71+0.874-222.981,0.004,0.75,0,0,1)
                        arcopreenchido(surface,137.97+40.723+279.5,556.31+0.874-222.981,6,1.75*(22/7),0.75*(22/7),137.97+40.723+279.5,556.31+0.874-222.981,0.004,0.75,0,0,1)
                        linhas(surface,0.75,0,0,1,1.214,133.74+40.723+279.5,560.47+0.874-222.981,150.72+40.723+279.5,543.4+0.874-222.981)
                        #linha tracejada religador de cima
                        linhastracejadas(surface,0.75,0,0,1,1.214,1,3,461.61,300,461.61,259.81)
                        escrita(surface,0.75, 0, 0,8.06,"sans-serif",0,0,360.365, 246.655,"Substituir por religador trifásico automático com")
                        escrita(surface,0.75, 0, 0,8.06,"sans-serif",0,0,360.365, 256.655,"sensor de presença de tensão em ambos os lados.")
            
                    elif dicionario_GUI['EquipContrucaoRede'] == 'Transformador' and PotUsinaModulo<300:
                        #escrita
                        escrita(surface,0, 0, 1,9.32,"sans-serif",0,0,466.387,293.213,f"TD {dicionario_GUI['NumEquipContrucaoRede']}")
                        #triangulo do trafo azul
                        linhas(surface,0,0,1,1,1.214,482.461,316.095,482.461,328.871)
                        triangulopreenchido(surface,473.854,316.095,491.068,316.095,482.461,300.095,1,1,1,0,0,1,1)
            
                    elif dicionario_GUI['EquipContrucaoRede'] == 'Fusível' or dicionario_GUI['EquipContrucaoRede'] == 'Chave' and PotUsinaModulo<300:
                        #escrita
                        escrita(surface,0.75, 0, 0,6.06,"Arial",0,1,449.330,527.562-218.151,f"FU {dicionario_GUI['NumEquipContrucaoRede']}")
                        #Fusivel vermelho
                        arcopreenchido(surface,146.5+40.723+279.5,547.71+0.874-222.981,6,0.75*(22/7),1.75*(22/7),146.5+40.723+279.5,547.71+0.874-222.981,0.004,0.75,0,0,1)
                        arcopreenchido(surface,137.97+40.723+279.5,556.31+0.874-222.981,6,1.75*(22/7),0.75*(22/7),137.97+40.723+279.5,556.31+0.874-222.981,0.004,0.75,0,0,1)
                        linhas(surface,0.75,0,0,1,1.214,133.74+40.723+279.5,560.47+0.874-222.981,150.72+40.723+279.5,543.4+0.874-222.981)
                        #linha tracejada religador de cima
                        linhastracejadas(surface,0.75,0,0,1,1.214,1,3,461.61,300,461.61,259.81)
                        escrita(surface,0.75, 0, 0,8.06,"sans-serif",0,0,416, 256.655,"Substituir por Chave Faca")
             
                x=0
                count = 0
                print('DESENHO RELIGADOR')
                print(Religador)
                for RL in Religador:
                    #Religador Vermelho
                    count += 1
                    losango(surface,31.65+54.172+x,547.58-218.151,41.38+54.172+x,537.85-218.151,51.11+54.172+x,547.58-218.151,41.38+54.172+x,557.32-218.151,1,1,1,0.75,0,0,1.8)
                    tamanho = 1.5*len(str(RL))*1.8
                    xlocal = (-float(31.65+54.172+x)+float(51.11+54.172+x))/2
                    xlocal = xlocal + (31.65+54.172+x) - tamanho
                    escrita(surface,0.75, 0, 0,6.06,"Arial",0,1,xlocal,527.562-218.151,f'RL {RL}')
                    escrita(surface,0.75, 0, 0,6.06,"Arial",0,1,xlocal,527.562-218.151-10,f'*({count})')
                    escrita(surface,0.75, 0, 0,8.06,"sans-serif",0,0,13,426+20*count,f"*({count})  Substituir por religador trifásico automático com sensor de presença de tensão em ambos os lados.")
                    x=x+distanciamento
                    
                for Re in Regulador:
                    #Regulador Vermelho
                    count += 1
                    arcopreenchido(surface,45.538+54.172+x,498.424-218.151+49.302,8.179,0,2*(22/7),45.538+54.172+x,498.424-218.151+49.302,0.003,0.75, 0, 0, 1)
                    tamanho = 4*len(str(Re))*1.8
                    xlocal = float(45.538+54.172+x)
                    xlocal = xlocal - tamanho
                    linhas(surface,0.75,0,0,1,1.214,33.686-0.521+54.172+x,511.954-0.954-218.151+49.302,60.134-0.521+54.172+x,485.506-0.954-218.151+49.302)
                    triangulo(surface,58.41-0.521+54.172+x,483.76-0.954-218.151+49.302,62.01-0.521+54.172+x,487.46-0.954-218.151+49.302,63.08-0.521+54.172+x,482.81-0.954-218.151+49.302,0.75,0,0)
                    escrita(surface,0.75, 0, 0,6.06,"Arial",0,1,xlocal,477.053-218.151+51,f"Regulador {Re}")
                    escrita(surface,0.75, 0, 0,6.06,"Arial",0,1,xlocal,477.053-218.151+51-10,f'*({count})')
                    escrita(surface,0.75, 0, 0,8.06,"sans-serif",0,0,13,426+20*count,f"*({count})  Substituir por regulador com funcionalidade de operação com fluxo reverso.")
                    x=x+distanciamento
                    
                for FUr in Fusivelreligador:
                    #Fusivel Vermelho
                    count += 1
                    arcopreenchido(surface,146.5+40.723+54.172-135.099+x,547.71+0.874-218.151+49.302-54.625,6,0.75*(22/7),1.75*(22/7),146.5+40.723+54.172-135.099+x,547.71+0.874-218.151+49.302-54.625,0.004,0.75,0,0,1)
                    arcopreenchido(surface,137.97+40.723+54.172-135.099+x,556.31+0.874-218.151+49.302-54.625,6,1.75*(22/7),0.75*(22/7),137.97+40.723+54.172-135.099+x,556.31+0.874-218.151+49.302-54.625,0.004,0.75,0,0,1)
                    linhas(surface,0.75,0,0,1,1.214,133.74+40.723+54.172-135.099+x,560.47+0.874-218.151+49.302-54.625,150.72+40.723+54.172-135.099+x,543.4+0.874-218.151+49.302-54.625)
                    tamanho = 2*len(str(FUr))*1.3
                    xlocal = (-float(133.74+40.723+54.172-135.099+x)+float(150.72+40.723+54.172-135.099+x))/2
                    xlocal = xlocal + (133.74+40.723+54.172-135.099+x) - tamanho
                    escrita(surface,0.75,0,0,6.06,"Arial",0,1,xlocal, 536.858-218.151+49.302-57.14,f"FU {FUr}")
                    escrita(surface,0.75,0,0,6.06,"Arial",0,1,xlocal, 536.858-218.151+49.302-57.14-10,f'*({count})')
                    escrita(surface,0.75,0,0,8.06,"sans-serif",0,0,13,426+20*count,f"*({count})  Substituir por religador trifásico automático com sensor de presença de tensão em ambos os lados.")
                    x=x+distanciamento
            
                for FUf in Fusivelfaca:
                    #Fusivel Vermelho
                    count += 1
                    arcopreenchido(surface,146.5+40.723+54.172-135.099+x,547.71+0.874-218.151+49.302-54.625,6,0.75*(22/7),1.75*(22/7),146.5+40.723+54.172-135.099+x,547.71+0.874-218.151+49.302-54.625,0.004,0.75,0,0,1)
                    arcopreenchido(surface,137.97+40.723+54.172-135.099+x,556.31+0.874-218.151+49.302-54.625,6,1.75*(22/7),0.75*(22/7),137.97+40.723+54.172-135.099+x,556.31+0.874-218.151+49.302-54.625,0.004,0.75,0,0,1)
                    linhas(surface,0.75,0,0,1,1.214,133.74+40.723+54.172-135.099+x,560.47+0.874-218.151+49.302-54.625,150.72+40.723+54.172-135.099+x,543.4+0.874-218.151+49.302-54.625)
                    tamanho = 2*len(str(FUf))*1.3
                    xlocal = (-float(133.74+40.723+54.172-135.099+x)+float(150.72+40.723+54.172-135.099+x))/2
                    xlocal = xlocal + (133.74+40.723+54.172-135.099+x) - tamanho
                    escrita(surface,0.75, 0, 0,6.06,"Arial",0,1,xlocal, 536.858-218.151+49.302-57.14,f"FU {FUf}")
                    escrita(surface,0.75, 0, 0,6.06,"Arial",0,1,xlocal, 536.858-218.151+49.302-57.14-10,f'*({count})')
                    escrita(surface,0.75, 0, 0,8.06,"sans-serif",0,0,13,426+20*count,f"*({count})  Substituir por Chave faca")
                    x=x+distanciamento
            
            return dicionario_word        
            
        dicionuro = {
            'Dia': '03',
            'Mes': 'Janeiro',
            'Ano': '2023',
            'Tratamento': '',
            'RepresentanteLegal': 'João José Medeiros de Figueiredo',
            'Endereco': 'R. Delfino Conti',
            'NumEndereco': 's/n',
            'CEP': '88040-900',
            'CidadeEndereco': 'Florianópolis',
            'EstadoEndereco': 'SC',
            'PrezadoTratamento': '',
            'Tipo': 'UFV',
            'NomeUsina': '',
            'NumSO': SO_automatica,
            'PotUsina': str(potencia_conexao*1000),
            'CidadeUsina': 'Florianópolis',
            'Alimentador': '',
            'Tensao': '',
            'ContrucaoRede': 'm',
            'MM': '',
            'TensaoCabo': '',
            'EquipContrucaoRede': '',
            'NumEquipContrucaoRede': '',
            'RecondMono': 'm',
            'TipoPontoAMono': '',
            'PontoAMono': '',
            'TipoPontoBMono': '',
            'PontoBMono': '',
            'Recond': 'm',
            'TipoPontoA': '',
            'PontoA': '',
            'TipoPontoB': '',
            'PontoB': '',
            'ReligadorSE': '',
            'Subestação': '',
            'Religadores': '',
            'kVBanco': '',
            'CorrenteBanco': '',
            'TipoPontoBanco': '',
            'NumPontoBanco': '',
            'Regulador': '',
            'FusiveisReligadores': '',
            'FusiveisFaca': '',
            'UTM_X': UTMX,
            'UTM_Y': UTMY,
            'ERDFinal': potencia_conexao*330.82}
        print('dicionario_GUI')
        print(dicionario_GUI)
        print('download_dir')
        print(download_dir)
        print('dicionario_obras_necessarias')
        print(dicionario_obras_necessarias)
        print('dicionario_quantidades')
        print(dicionario_quantidades)
        dicionario_word = Construcao_dicionarios(dicionuro,dicionario_GUI)
        print('dicionario_word')
        print(dicionario_word)
        caminho_pasta = "C:\\Users\\joaof\\OneDrive\\Área de Trabalho\\TCC\\programa atualizado\\Words_testes"
        a_o(caminho_pasta,dicionario_word,download_dir,dicionario_obras_necessarias,dicionario_quantidades,dicionario_GUI)
    PEPtoWORD(numero_so,caminhoSO,dicionario_obras_necessarias,dicionario_quantidades,dicionario_GUI)
